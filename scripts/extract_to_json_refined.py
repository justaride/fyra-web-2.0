"""
Refined extraction with better data cleaning and structure
"""

import json
from docx import Document
import re
from pathlib import Path

def clean_text(text):
    """Clean and normalize text"""
    if not text:
        return ""
    return ' '.join(text.strip().split())

def is_header(para):
    """Check if paragraph is a header"""
    if not para.text.strip():
        return False
    if para.style.name.startswith('Heading'):
        return True
    if para.runs and para.runs[0].bold and len(para.text) < 150:
        return True
    return False

def extract_suppliers_refined(doc_paths):
    """Extract B2B suppliers with better structure"""
    suppliers = []

    for doc_path in doc_paths:
        doc = Document(doc_path)
        current_supplier = None
        current_section = None

        for para in doc.paragraphs:
            text = clean_text(para.text)
            if not text or len(text) < 3:
                continue

            # Skip document headers and tier classifications
            if any(skip in text.upper() for skip in ['COMPREHENSIVE B2B', 'TIER 1:', 'TIER 2:', 'TIER 3:',
                                                       'CONTRACT-LEVEL', 'SVERIGE -', 'REGION']):
                continue

            # Detect numbered supplier entries (1. NAME, 2. NAME, etc.)
            if re.match(r'^\d+\.\s+[A-ZÅÄÖ]', text):
                if current_supplier and current_supplier.get('name'):
                    suppliers.append(current_supplier)

                # Extract name from numbered format
                name = re.sub(r'^\d+\.\s+', '', text)
                name = re.sub(r'\s*\(.*?\)\s*', '', name)  # Remove parenthetical notes
                current_supplier = {
                    'name': name.strip(),
                    'type': 'B2B Reuse Operator',
                    'services': [],
                    'capabilities': [],
                    'certifications': []
                }
                current_section = None

            elif current_supplier:
                # Parse structured data
                if ':' in text:
                    parts = text.split(':', 1)
                    key = parts[0].strip().lower()
                    value = parts[1].strip() if len(parts) > 1 else ""

                    if not value or len(value) < 2:
                        continue

                    # Map fields
                    if any(k in key for k in ['location', 'lokalisering', 'stad', 'city']):
                        current_supplier['location'] = value
                    elif 'website' in key or 'webbplats' in key:
                        current_supplier['website'] = value
                    elif 'email' in key or 'e-post' in key:
                        current_supplier['email'] = value
                    elif 'phone' in key or 'telefon' in key:
                        current_supplier['phone'] = value
                    elif 'contact' in key or 'kontakt' in key:
                        current_supplier['contact'] = value
                    elif any(k in key for k in ['focus', 'specialty', 'specialitet']):
                        current_supplier['specialty'] = value
                    elif 'volume' in key or 'volym' in key:
                        current_supplier['volumeCapacity'] = value
                    elif 'quality' in key or 'kvalitet' in key:
                        current_supplier['qualityAssurance'] = value
                    elif 'hospitality readiness' in key:
                        current_supplier['hospitalityReadiness'] = value
                    elif 'best for' in key:
                        current_supplier['bestFor'] = value
                    elif 'key strength' in key:
                        current_supplier['keyStrength'] = value
                    elif 'key gap' in key:
                        current_supplier['keyGap'] = value

        # Add last supplier
        if current_supplier and current_supplier.get('name'):
            suppliers.append(current_supplier)

    # Clean and deduplicate
    unique = {}
    for s in suppliers:
        name = s.get('name', '').strip()
        if name and len(name) > 2 and not any(skip in name.upper() for skip in
            ['CRITICAL', 'NOTE:', 'TIER', 'HOSPITALITY', 'FOCUS:', 'KEY']):
            if name not in unique:
                unique[name] = s
            else:
                # Merge data
                for k, v in s.items():
                    if k not in unique[name] or not unique[name][k]:
                        unique[name][k] = v

    return list(unique.values())

def extract_consultants_refined(doc_path):
    """Extract consultants with better structure"""
    doc = Document(doc_path)
    consultants = []
    current_consultant = None

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text or len(text) < 3:
            continue

        # Skip document headers
        if any(skip in text.upper() for skip in ['SWEDISH PROJECT MANAGERS', 'TIER 1:', 'TIER 2:',
                                                   'EXECUTIVE SUMMARY', 'IMMEDIATE ACTION']):
            continue

        # Detect numbered consultant entries
        if re.match(r'^\d+\.\s+[A-ZÅÄÖ]', text):
            if current_consultant and current_consultant.get('name'):
                consultants.append(current_consultant)

            name = re.sub(r'^\d+\.\s+', '', text)
            name = re.sub(r'\s*-\s*PRIORITY.*', '', name)  # Remove priority notes
            current_consultant = {
                'name': name.strip(),
                'type': 'Project Manager / Technical Consultant',
                'projects': [],
                'certifications': [],
                'expertise': []
            }

        elif current_consultant:
            if ':' in text:
                parts = text.split(':', 1)
                key = parts[0].strip().lower()
                value = parts[1].strip() if len(parts) > 1 else ""

                if not value or len(value) < 2:
                    continue

                if 'location' in key or 'stad' in key:
                    current_consultant['location'] = value
                elif 'website' in key:
                    current_consultant['website'] = value
                elif 'contact' in key or 'email' in key:
                    current_consultant['contact'] = value
                elif 'phone' in key or 'telefon' in key:
                    current_consultant['phone'] = value
                elif 'certification' in key:
                    if value not in current_consultant['certifications']:
                        current_consultant['certifications'].append(value)
                elif 'hospitality track record' in key:
                    current_consultant['hospitalityExperience'] = value
                elif 'circular economy' in key:
                    current_consultant['circularEconomyExperience'] = value
                elif 'key strength' in key:
                    current_consultant['keyStrength'] = value

    if current_consultant and current_consultant.get('name'):
        consultants.append(current_consultant)

    # Clean
    result = []
    for c in consultants:
        name = c.get('name', '').strip()
        if name and len(name) > 2 and not any(skip in name.upper() for skip in
            ['CONTACT &', 'KEY CONTACTS', 'HOSPITALITY']):
            result.append(c)

    return result

def extract_case_studies_refined(doc_path):
    """Extract hotel case studies with better structure"""
    doc = Document(doc_path)
    case_studies = []
    current_case = None

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text or len(text) < 3:
            continue

        # Skip document headers
        if any(skip in text.upper() for skip in ['FRONTRUNNER HOTELS', 'TIER 1:', 'PRIORITY TIER']):
            continue

        # Detect hotel names (typically all caps or bold)
        if is_header(para) and not ':' in text:
            if re.match(r'^[A-ZÅÄÖ\s&-]+$', text) or 'HOTEL' in text.upper() or 'BY NOBIS' in text.upper():
                if current_case and current_case.get('name'):
                    case_studies.append(current_case)

                # Clean hotel name
                name = re.sub(r'\s*-\s*RELEVANCE SCORE:.*', '', text)
                current_case = {
                    'name': name.strip(),
                    'type': 'Hotel Case Study',
                    'circularElements': [],
                    'materials': []
                }

        elif current_case:
            if ':' in text:
                parts = text.split(':', 1)
                key = parts[0].strip().lower()
                value = parts[1].strip() if len(parts) > 1 else ""

                if not value or len(value) < 2:
                    continue

                if 'name:' in text.lower() and 'hotel' in value.lower():
                    current_case['name'] = value
                elif 'location' in key or 'address' in key:
                    current_case['location'] = value
                elif 'size' in key or 'rooms' in key:
                    current_case['size'] = value
                elif 'opening' in key or 'year' in key:
                    current_case['year'] = value
                elif 'website' in key:
                    current_case['website'] = value
                elif 'phone' in key:
                    current_case['phone'] = value
                elif 'category' in key or 'type' in key:
                    current_case['category'] = value
                elif 'chain' in key or 'brand' in key:
                    current_case['chain'] = value
                elif 'scope' in key:
                    current_case['scope'] = value
                elif 'co2' in key.lower() or 'carbon' in key:
                    current_case['co2Savings'] = value
                elif 'financial' in key or 'cost' in key or 'budget' in key:
                    current_case['financialImpact'] = value
                elif 'waste' in key:
                    current_case['wasteReduction'] = value
                elif 'architect' in key:
                    current_case['architect'] = value
                elif 'contractor' in key:
                    current_case['contractor'] = value

    if current_case and current_case.get('name'):
        case_studies.append(current_case)

    # Clean
    result = []
    for c in case_studies:
        name = c.get('name', '').strip()
        if name and len(name) > 2 and not any(skip in name.upper() for skip in
            ['PROFILE', 'CIRCULAR ECONOMY', 'QUANTIFIED', 'KEY PEOPLE']):
            result.append(c)

    return result

def extract_regulations_refined(doc_path):
    """Extract regulations with better categorization"""
    doc = Document(doc_path)

    regulations = {
        'fireSafety': {'title': 'Fire Safety Requirements', 'items': []},
        'buildingCodes': {'title': 'Building Codes (BBR)', 'items': []},
        'bvb': {'title': 'BVB Standards', 'items': []},
        'materialStandards': {'title': 'Material Standards', 'items': []},
        'documentation': {'title': 'Documentation Requirements', 'items': []},
        'compliance': {'title': 'Compliance & Certification', 'items': []}
    }

    current_category = 'documentation'
    current_item = None

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text or len(text) < 3:
            continue

        # Detect category changes
        text_lower = text.lower()
        if 'fire safety' in text_lower or 'brand' in text_lower:
            current_category = 'fireSafety'
        elif 'bbr' in text_lower or 'building code' in text_lower or 'boverket' in text_lower:
            current_category = 'buildingCodes'
        elif 'bvb' in text_lower and 'standard' in text_lower:
            current_category = 'bvb'
        elif 'material standard' in text_lower or 'ce marking' in text_lower:
            current_category = 'materialStandards'
        elif 'documentation' in text_lower or 'dokumentation' in text_lower:
            current_category = 'documentation'
        elif 'compliance' in text_lower or 'certification' in text_lower:
            current_category = 'compliance'

        # Add items
        if is_header(para):
            if current_item:
                regulations[current_category]['items'].append(current_item)
            current_item = {'title': text, 'description': ''}
        elif current_item:
            current_item['description'] += ' ' + text
        else:
            regulations[current_category]['items'].append({'title': text, 'description': ''})

    if current_item:
        regulations[current_category]['items'].append(current_item)

    # Clean up
    for category in regulations:
        regulations[category]['items'] = [item for item in regulations[category]['items']
                                           if item.get('title') and len(item['title']) > 10]

    return regulations

def extract_implementation_refined(doc_path):
    """Extract implementation roadmap with better structure"""
    doc = Document(doc_path)

    implementation = {
        'executiveSummary': '',
        'marketMaturity': '',
        'strengths': [],
        'weaknesses': [],
        'roadmap': {
            'phase1': {'title': 'Phase 1', 'timeline': '', 'actions': []},
            'phase2': {'title': 'Phase 2', 'timeline': '', 'actions': []},
            'phase3': {'title': 'Phase 3', 'timeline': '', 'actions': []}
        },
        'marketAnalysis': {
            'competitors': [],
            'opportunities': [],
            'threats': []
        },
        'recommendations': [],
        'kpis': []
    }

    current_section = 'executiveSummary'

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text or len(text) < 3:
            continue

        text_lower = text.lower()

        # Detect sections
        if 'executive summary' in text_lower:
            current_section = 'executiveSummary'
        elif 'market maturity' in text_lower or 'readiness' in text_lower:
            current_section = 'marketMaturity'
        elif 'strength' in text_lower:
            current_section = 'strengths'
        elif 'weakness' in text_lower or 'gap' in text_lower:
            current_section = 'weaknesses'
        elif 'recommendation' in text_lower:
            current_section = 'recommendations'
        elif 'kpi' in text_lower or 'metric' in text_lower:
            current_section = 'kpis'

        # Store content
        if current_section == 'executiveSummary':
            if implementation['executiveSummary']:
                implementation['executiveSummary'] += ' ' + text
            else:
                implementation['executiveSummary'] = text
        elif current_section == 'marketMaturity':
            if implementation['marketMaturity']:
                implementation['marketMaturity'] += ' ' + text
            else:
                implementation['marketMaturity'] = text
        elif current_section in ['strengths', 'weaknesses', 'recommendations', 'kpis']:
            if text and len(text) > 10:
                implementation[current_section].append(text)

    return implementation

def main():
    """Main extraction process"""
    base_path = Path('/Users/gabrielboen/Downloads/drive-download-20251108T110451Z-1-001')

    print("Starting refined data extraction...")

    # Extract suppliers
    print("\n1. Extracting B2B Reuse Operators (refined)...")
    suppliers = extract_suppliers_refined([
        base_path / '1. COMPREHENSIVE B2B REUSE OPERATORS EXTRACTION.docx',
        base_path / '2. FYRA B2B OPERATOR CAPABILITY MATRIX.docx'
    ])

    with open(base_path / 'suppliers.json', 'w', encoding='utf-8') as f:
        json.dump(suppliers, f, indent=2, ensure_ascii=False)
    print(f"   ✓ {len(suppliers)} suppliers → suppliers.json")

    # Extract consultants
    print("\n2. Extracting Consultants (refined)...")
    consultants = extract_consultants_refined(
        base_path / '3. SWEDISH PROJECT MANAGERS & TECHNICAL CONSULTANTS FOR WHOLE-PROJECT HOSPITALITY RENOVATION.docx'
    )

    with open(base_path / 'consultants.json', 'w', encoding='utf-8') as f:
        json.dump(consultants, f, indent=2, ensure_ascii=False)
    print(f"   ✓ {len(consultants)} consultants → consultants.json")

    # Extract case studies
    print("\n3. Extracting Hotel Case Studies (refined)...")
    case_studies = extract_case_studies_refined(
        base_path / '4. FRONTRUNNER HOTELS - COMPLETE EXTRACTION.docx'
    )

    with open(base_path / 'caseStudies.json', 'w', encoding='utf-8') as f:
        json.dump(case_studies, f, indent=2, ensure_ascii=False)
    print(f"   ✓ {len(case_studies)} case studies → caseStudies.json")

    # Extract regulations
    print("\n4. Extracting Regulations (refined)...")
    regulations = extract_regulations_refined(
        base_path / '5. PRACTICAL REGULATORY GUIDE_ REUSED MATERIALS IN SWEDISH HOTEL RENOVATIONS.docx'
    )

    with open(base_path / 'regulations.json', 'w', encoding='utf-8') as f:
        json.dump(regulations, f, indent=2, ensure_ascii=False)

    total_regs = sum(len(v['items']) for v in regulations.values())
    print(f"   ✓ {total_regs} regulations → regulations.json")

    # Extract implementation
    print("\n5. Extracting Implementation Roadmap (refined)...")
    implementation = extract_implementation_refined(
        base_path / '6. STRATEGIC IMPLEMENTATION ROADMAP.docx'
    )

    with open(base_path / 'implementation.json', 'w', encoding='utf-8') as f:
        json.dump(implementation, f, indent=2, ensure_ascii=False)
    print(f"   ✓ Implementation roadmap → implementation.json")

    # Create summary
    summary = {
        'extractionDate': '2025-11-08',
        'version': '2.0-refined',
        'sourceDocuments': 6,
        'outputFiles': 5,
        'dataExtracted': {
            'suppliers': len(suppliers),
            'consultants': len(consultants),
            'caseStudies': len(case_studies),
            'regulations': total_regs,
            'regulationCategories': len(regulations)
        },
        'files': {
            'suppliers': 'suppliers.json - B2B reuse operators with contact info, capabilities, and specialty areas',
            'consultants': 'consultants.json - Project managers and technical consultants for hotel renovations',
            'caseStudies': 'caseStudies.json - Real hotel projects with circular economy implementation',
            'regulations': 'regulations.json - Swedish regulatory requirements organized by category',
            'implementation': 'implementation.json - Strategic roadmap and market analysis'
        },
        'dataStructure': {
            'suppliers': {
                'fields': ['name', 'type', 'location', 'website', 'contact', 'specialty', 'volumeCapacity',
                          'qualityAssurance', 'hospitalityReadiness', 'bestFor', 'keyStrength'],
                'sampleCount': min(3, len(suppliers))
            },
            'consultants': {
                'fields': ['name', 'type', 'location', 'website', 'contact', 'certifications',
                          'hospitalityExperience', 'circularEconomyExperience', 'keyStrength'],
                'sampleCount': min(3, len(consultants))
            },
            'caseStudies': {
                'fields': ['name', 'type', 'location', 'size', 'year', 'chain', 'scope',
                          'co2Savings', 'financialImpact', 'architect', 'contractor'],
                'sampleCount': min(3, len(case_studies))
            },
            'regulations': {
                'categories': list(regulations.keys()),
                'structure': 'Each category contains title and array of items with title/description'
            }
        }
    }

    with open(base_path / 'extraction_summary.json', 'w', encoding='utf-8') as f:
        json.dump(summary, f, indent=2, ensure_ascii=False)

    print("\n" + "="*70)
    print("REFINED EXTRACTION COMPLETE!")
    print("="*70)
    print(f"\nData extracted:")
    print(f"  • {len(suppliers)} B2B Reuse Operators")
    print(f"  • {len(consultants)} Project Managers/Consultants")
    print(f"  • {len(case_studies)} Hotel Case Studies")
    print(f"  • {total_regs} Regulatory Requirements across {len(regulations)} categories")
    print(f"  • Strategic Implementation Roadmap")
    print(f"\nJSON files created in: {base_path}")
    print(f"  1. suppliers.json")
    print(f"  2. consultants.json")
    print(f"  3. caseStudies.json")
    print(f"  4. regulations.json")
    print(f"  5. implementation.json")
    print(f"  6. extraction_summary.json")
    print("\nReady for Next.js integration!")

if __name__ == '__main__':
    main()
