#!/usr/bin/env python3
"""
Properly extract data from Fyra documents based on actual document structure.
"""

import json
import re
from pathlib import Path
from docx import Document

BASE_DIR = Path("/Users/gabrielboen/Downloads/drive-download-20251108T110451Z-1-001")
DATA_DIR = BASE_DIR / "fyra-circular-platform" / "data"

def clean_text(text):
    """Clean text."""
    if not text:
        return ""
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)
    return text

def extract_contact_info(text):
    """Extract contact info."""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = list(set(re.findall(email_pattern, text)))

    phone_pattern = r'\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}'
    phones = re.findall(phone_pattern, text)
    phones = [p for p in phones if len(p.replace('-', '').replace(' ', '').replace('.', '')) >= 8]

    website_pattern = r'(?:https?://)?(?:www\.)?[a-zA-Z0-9-]+\.[a-zA-Z]{2,}(?:/[^\s]*)?'
    websites = [w for w in re.findall(website_pattern, text) if '@' not in w]

    return {
        "emails": emails[:1] if emails else [],
        "phones": phones[:1] if phones else [],
        "websites": websites[:1] if websites else []
    }

def parse_document_sections(doc):
    """Parse document into hierarchical sections."""
    sections = []
    current_h1 = None
    current_h2 = None
    current_h3 = None

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        style = para.style.name

        if style == 'Heading 1':
            if current_h1:
                sections.append(current_h1)
            current_h1 = {'title': text, 'h2': [], 'content': []}
            current_h2 = None
            current_h3 = None

        elif style == 'Heading 2':
            if current_h1:
                if current_h2:
                    current_h1['h2'].append(current_h2)
                current_h2 = {'title': text, 'h3': [], 'content': []}
                current_h3 = None

        elif style == 'Heading 3':
            if current_h2:
                if current_h3:
                    current_h2['h3'].append(current_h3)
                current_h3 = {'title': text, 'content': []}

        else:  # Normal text
            if current_h3:
                current_h3['content'].append(text)
            elif current_h2:
                current_h2['content'].append(text)
            elif current_h1:
                current_h1['content'].append(text)

    # Add final sections
    if current_h3 and current_h2:
        current_h2['h3'].append(current_h3)
    if current_h2 and current_h1:
        current_h1['h2'].append(current_h2)
    if current_h1:
        sections.append(current_h1)

    return sections

def extract_suppliers():
    """Extract B2B suppliers."""
    doc_path = BASE_DIR / "1. COMPREHENSIVE B2B REUSE OPERATORS EXTRACTION.docx"
    doc = Document(doc_path)
    sections = parse_document_sections(doc)

    suppliers = []

    for h1 in sections:
        for h2 in h1.get('h2', []):
            for h3 in h2.get('h3', []):
                name = h3['title']

                # Skip non-company sections
                if any(skip in name.lower() for skip in ['tier', 'summary', 'note', 'action']):
                    continue

                # Clean company name (remove numbers like "1. ", "2. ")
                name = re.sub(r'^\d+\.\s*', '', name)

                content = '\n'.join(h3['content'])
                contacts = extract_contact_info(content)

                supplier = {
                    "id": name.lower().replace(' ', '_').replace('å', 'a').replace('ä', 'a').replace('ö', 'o').replace('/', '_')[:50],
                    "name": name,
                    "description": "",
                    "location": "",
                    "services": [],
                    "capabilities": {
                        "volume": "",
                        "leadTime": "",
                        "inventory": "",
                        "logistics": ""
                    },
                    "contact": {
                        "name": "",
                        "phone": contacts["phones"][0] if contacts["phones"] else "",
                        "email": contacts["emails"][0] if contacts["emails"] else "",
                        "website": contacts["websites"][0] if contacts["websites"] else ""
                    },
                    "certifications": [],
                    "hospitalityReadiness": {
                        "score": "",
                        "strengths": [],
                        "gaps": []
                    },
                    "pricing": "",
                    "projectExamples": []
                }

                # Parse content
                for line in h3['content']:
                    lower = line.lower()

                    if line.startswith('Location:'):
                        supplier["location"] = line.split(':', 1)[1].strip()
                    elif line.startswith('Focus:') or line.startswith('Services:'):
                        supplier["description"] = line.split(':', 1)[1].strip()
                    elif 'hospitality readiness' in lower:
                        supplier["hospitalityReadiness"]["score"] = line
                    elif line.startswith('Key Strength') or line.startswith('Strengths:'):
                        supplier["hospitalityReadiness"]["strengths"].append(line.split(':', 1)[1].strip() if ':' in line else line)
                    elif line.startswith('Key Gap') or line.startswith('Gaps:'):
                        supplier["hospitalityReadiness"]["gaps"].append(line.split(':', 1)[1].strip() if ':' in line else line)
                    elif line.startswith('Best For:') or line.startswith('Services:'):
                        services = line.split(':', 1)[1].strip() if ':' in line else line
                        supplier["services"].append(services)
                    elif 'volume' in lower or 'capacity' in lower:
                        supplier["capabilities"]["volume"] = line
                    elif 'lead time' in lower or 'delivery' in lower:
                        supplier["capabilities"]["leadTime"] = line

                if not supplier["description"] and h3['content']:
                    supplier["description"] = h3['content'][0] if h3['content'] else ""

                suppliers.append(supplier)

    return suppliers

def extract_consultants():
    """Extract consultants."""
    doc_path = BASE_DIR / "3. SWEDISH PROJECT MANAGERS & TECHNICAL CONSULTANTS FOR WHOLE-PROJECT HOSPITALITY RENOVATION.docx"
    doc = Document(doc_path)
    sections = parse_document_sections(doc)

    consultants = []

    for h1 in sections:
        for h2 in h1.get('h2', []):
            # H2 level might be company names
            for h3 in h2.get('h3', []):
                name = h3['title']

                # Skip non-company sections
                if any(skip in name.lower() for skip in ['tier', 'summary', 'action', 'contact', 'agenda', 'gap', 'strategy']):
                    continue

                name = re.sub(r'^\d+\.\s*', '', name)

                content = '\n'.join(h3['content'])
                contacts = extract_contact_info(content)

                consultant = {
                    "id": name.lower().replace(' ', '_').replace('å', 'a').replace('ä', 'a').replace('ö', 'o').replace('/', '_')[:50],
                    "name": name,
                    "description": "",
                    "services": [],
                    "specializations": [],
                    "hospitalityExperience": {
                        "projects": [],
                        "expertise": ""
                    },
                    "circularEconomyExpertise": "",
                    "contact": {
                        "name": "",
                        "phone": contacts["phones"][0] if contacts["phones"] else "",
                        "email": contacts["emails"][0] if contacts["emails"] else "",
                        "website": contacts["websites"][0] if contacts["websites"] else ""
                    },
                    "certifications": [],
                    "strengths": []
                }

                # Parse content
                for line in h3['content']:
                    lower = line.lower()

                    if 'business model' in lower:
                        consultant["description"] = line.split(':', 1)[1].strip() if ':' in line else line
                    elif 'hospitality' in lower and 'project' in lower:
                        consultant["hospitalityExperience"]["projects"].append(line)
                    elif 'circular' in lower or 'sustainability' in lower:
                        if consultant["circularEconomyExpertise"]:
                            consultant["circularEconomyExpertise"] += " " + line
                        else:
                            consultant["circularEconomyExpertise"] = line
                    elif line.startswith('Services:'):
                        consultant["services"].append(line.split(':', 1)[1].strip())
                    elif 'certification' in lower or 'certified' in lower:
                        consultant["certifications"].append(line)
                    elif line.startswith('Key Strength') or 'strength' in lower:
                        consultant["strengths"].append(line)

                if not consultant["description"] and h3['content']:
                    consultant["description"] = h3['content'][0] if h3['content'] else ""

                consultants.append(consultant)

    return consultants

def extract_case_studies():
    """Extract hotel case studies."""
    doc_path = BASE_DIR / "4. FRONTRUNNER HOTELS - COMPLETE EXTRACTION.docx"
    doc = Document(doc_path)
    sections = parse_document_sections(doc)

    case_studies = []

    for h1 in sections:
        for h2 in h1.get('h2', []):
            for h3 in h2.get('h3', []):
                name = h3['title']

                # Skip non-hotel sections
                if any(skip in name.lower() for skip in ['introduction', 'summary', 'methodology', 'tier']):
                    continue

                # Hotel names are often in H3 or H2
                hotel_name = name

                content = '\n'.join(h3['content'])

                case_study = {
                    "id": hotel_name.lower().replace(' ', '_').replace('å', 'a').replace('ä', 'a').replace('ö', 'o')[:50],
                    "hotelName": hotel_name,
                    "location": h1['title'] if 'sverige' in h1['title'].lower() or 'denmark' in h1['title'].lower() or 'norway' in h1['title'].lower() else "",
                    "year": "",
                    "projectSize": "",
                    "category": "",
                    "circularElements": [],
                    "impact": {
                        "co2Savings": "",
                        "circularPercentage": "",
                        "costSavings": ""
                    },
                    "architects": [],
                    "partners": [],
                    "challenges": [],
                    "outcomes": [],
                    "relevance": ""
                }

                # Extract year from content
                year_match = re.search(r'\b(19|20)\d{2}\b', content)
                if year_match:
                    case_study["year"] = year_match.group()

                # Parse content
                for line in h3['content']:
                    lower = line.lower()

                    if 'location' in lower or 'city' in lower:
                        if ':' in line:
                            case_study["location"] = line.split(':', 1)[1].strip()
                    elif 'year' in lower or 'completed' in lower or 'opened' in lower:
                        if ':' in line:
                            case_study["year"] = line.split(':', 1)[1].strip()
                    elif 'circular' in lower or 'reused' in lower or 'reuse' in lower or 'material' in lower:
                        case_study["circularElements"].append(line)
                    elif 'co2' in lower or 'carbon' in lower:
                        case_study["impact"]["co2Savings"] = line
                    elif '%' in line and 'circular' in lower:
                        case_study["impact"]["circularPercentage"] = line
                    elif 'architect' in lower:
                        case_study["architects"].append(line)
                    elif 'sustainability consultant' in lower or 'material supplier' in lower:
                        case_study["partners"].append(line)
                    elif 'challenge' in lower:
                        case_study["challenges"].append(line)
                    elif 'outcome' in lower or 'result' in lower:
                        case_study["outcomes"].append(line)
                    elif 'justification' in lower or 'relevance' in lower:
                        case_study["relevance"] = line

                if case_study["circularElements"] or case_study["impact"]["co2Savings"]:
                    case_studies.append(case_study)

    return case_studies

def extract_regulatory():
    """Extract regulatory info."""
    doc_path = BASE_DIR / "5. PRACTICAL REGULATORY GUIDE_ REUSED MATERIALS IN SWEDISH HOTEL RENOVATIONS.docx"
    doc = Document(doc_path)
    sections = parse_document_sections(doc)

    regulatory = {
        "fireSafety": {
            "overview": "",
            "bbrRequirements": [],
            "testingStandards": [],
            "documentation": []
        },
        "buildingCodes": {
            "overview": "",
            "hotelSpecific": [],
            "materialRequirements": []
        },
        "bvbStandards": {
            "overview": "",
            "scoring": [],
            "certificationProcess": []
        },
        "documentation": {
            "required": [],
            "recommended": [],
            "templates": []
        },
        "implementation": {
            "tips": [],
            "bestPractices": [],
            "commonChallenges": []
        }
    }

    for h1 in sections:
        title_lower = h1['title'].lower()

        if 'fire' in title_lower or 'brand' in title_lower:
            section = 'fireSafety'
        elif 'bbr' in title_lower or 'building code' in title_lower or 'byggregler' in title_lower:
            section = 'buildingCodes'
        elif 'bvb' in title_lower or 'environmental' in title_lower or 'miljö' in title_lower:
            section = 'bvbStandards'
        elif 'documentation' in title_lower or 'dokument' in title_lower:
            section = 'documentation'
        elif 'implementation' in title_lower or 'practical' in title_lower or 'praktisk' in title_lower:
            section = 'implementation'
        else:
            continue

        # Add overview
        if h1.get('content'):
            regulatory[section]["overview"] = ' '.join(h1['content'][:3])

        # Process H2 sections
        for h2 in h1.get('h2', []):
            for line in h2.get('content', []):
                if section == 'fireSafety':
                    if 'bbr' in line.lower() or 'requirement' in line.lower():
                        regulatory[section]["bbrRequirements"].append(line)
                    elif 'test' in line.lower() or 'standard' in line.lower():
                        regulatory[section]["testingStandards"].append(line)
                    else:
                        regulatory[section]["documentation"].append(line)

                elif section == 'buildingCodes':
                    if 'hotel' in line.lower():
                        regulatory[section]["hotelSpecific"].append(line)
                    else:
                        regulatory[section]["materialRequirements"].append(line)

                elif section == 'bvbStandards':
                    if 'point' in line.lower() or 'score' in line.lower() or 'rating' in line.lower():
                        regulatory[section]["scoring"].append(line)
                    else:
                        regulatory[section]["certificationProcess"].append(line)

                elif section == 'documentation':
                    if 'required' in line.lower() or 'must' in line.lower():
                        regulatory[section]["required"].append(line)
                    elif 'recommended' in line.lower() or 'should' in line.lower():
                        regulatory[section]["recommended"].append(line)
                    else:
                        regulatory[section]["templates"].append(line)

                elif section == 'implementation':
                    if 'tip' in line.lower() or 'advice' in line.lower():
                        regulatory[section]["tips"].append(line)
                    elif 'best practice' in line.lower() or 'recommendation' in line.lower():
                        regulatory[section]["bestPractices"].append(line)
                    elif 'challenge' in line.lower() or 'risk' in line.lower():
                        regulatory[section]["commonChallenges"].append(line)
                    else:
                        regulatory[section]["tips"].append(line)

    return regulatory

def extract_strategy():
    """Extract strategy."""
    doc_path = BASE_DIR / "6. STRATEGIC IMPLEMENTATION ROADMAP.docx"
    doc = Document(doc_path)
    sections = parse_document_sections(doc)

    strategy = {
        "marketAnalysis": {
            "swedenOverview": "",
            "hotelSector": "",
            "circularEconomyTrends": "",
            "opportunities": [],
            "challenges": []
        },
        "fyraPositioning": {
            "valueProposition": "",
            "differentiators": [],
            "targetMarket": "",
            "competitiveAdvantages": []
        },
        "recommendations": {
            "immediate": [],
            "shortTerm": [],
            "longTerm": []
        },
        "implementation": {
            "roadmap": [],
            "milestones": [],
            "resources": [],
            "successMetrics": []
        }
    }

    for h1 in sections:
        title_lower = h1['title'].lower()

        # Identify main section
        if 'market' in title_lower or 'analys' in title_lower:
            section = 'marketAnalysis'
        elif 'fyra' in title_lower or 'position' in title_lower:
            section = 'fyraPositioning'
        elif 'recommend' in title_lower:
            section = 'recommendations'
        elif 'implementation' in title_lower or 'roadmap' in title_lower:
            section = 'implementation'
        else:
            continue

        # Add overview text
        if section == 'marketAnalysis' and h1.get('content'):
            strategy[section]["swedenOverview"] = ' '.join(h1['content'][:5])
        elif section == 'fyraPositioning' and h1.get('content'):
            strategy[section]["valueProposition"] = ' '.join(h1['content'][:5])

        # Process H2 and content
        for h2 in h1.get('h2', []):
            h2_lower = h2['title'].lower()

            for line in h2.get('content', []):
                if section == 'marketAnalysis':
                    if 'opportunit' in h2_lower:
                        strategy[section]["opportunities"].append(line)
                    elif 'challenge' in h2_lower:
                        strategy[section]["challenges"].append(line)
                    elif 'hotel' in h2_lower:
                        strategy[section]["hotelSector"] += " " + line
                    elif 'circular' in h2_lower:
                        strategy[section]["circularEconomyTrends"] += " " + line

                elif section == 'fyraPositioning':
                    if 'different' in h2_lower or 'unique' in h2_lower:
                        strategy[section]["differentiators"].append(line)
                    elif 'advantage' in h2_lower or 'competitive' in h2_lower:
                        strategy[section]["competitiveAdvantages"].append(line)
                    elif 'target' in h2_lower:
                        strategy[section]["targetMarket"] += " " + line

                elif section == 'recommendations':
                    if 'immediate' in h2_lower or 'short' in h2_lower or 'quick' in h2_lower:
                        strategy[section]["immediate"].append(line)
                    elif 'medium' in h2_lower:
                        strategy[section]["shortTerm"].append(line)
                    elif 'long' in h2_lower:
                        strategy[section]["longTerm"].append(line)
                    else:
                        strategy[section]["immediate"].append(line)

                elif section == 'implementation':
                    if 'roadmap' in h2_lower or 'timeline' in h2_lower or 'phase' in h2_lower:
                        strategy[section]["roadmap"].append(line)
                    elif 'milestone' in h2_lower:
                        strategy[section]["milestones"].append(line)
                    elif 'resource' in h2_lower:
                        strategy[section]["resources"].append(line)
                    elif 'metric' in h2_lower or 'kpi' in h2_lower or 'success' in h2_lower:
                        strategy[section]["successMetrics"].append(line)
                    else:
                        strategy[section]["roadmap"].append(line)

    # Clean up text fields
    strategy["marketAnalysis"]["swedenOverview"] = clean_text(strategy["marketAnalysis"]["swedenOverview"])
    strategy["marketAnalysis"]["hotelSector"] = clean_text(strategy["marketAnalysis"]["hotelSector"])
    strategy["marketAnalysis"]["circularEconomyTrends"] = clean_text(strategy["marketAnalysis"]["circularEconomyTrends"])
    strategy["fyraPositioning"]["valueProposition"] = clean_text(strategy["fyraPositioning"]["valueProposition"])
    strategy["fyraPositioning"]["targetMarket"] = clean_text(strategy["fyraPositioning"]["targetMarket"])

    return strategy

def main():
    print("=" * 80)
    print("FYRA CIRCULAR PLATFORM - COMPREHENSIVE DATA EXTRACTION")
    print("=" * 80)

    print("\n[1/5] Extracting suppliers...")
    suppliers = extract_suppliers()
    print(f"      ✓ Extracted {len(suppliers)} suppliers")

    print("\n[2/5] Extracting consultants...")
    consultants = extract_consultants()
    print(f"      ✓ Extracted {len(consultants)} consultants")

    print("\n[3/5] Extracting case studies...")
    case_studies = extract_case_studies()
    print(f"      ✓ Extracted {len(case_studies)} case studies")

    print("\n[4/5] Extracting regulatory info...")
    regulatory = extract_regulatory()
    total_reg = sum(len(v) if isinstance(v, list) else 0 for section in regulatory.values() for v in section.values())
    print(f"      ✓ Extracted {total_reg} regulatory items")

    print("\n[5/5] Extracting strategy...")
    strategy = extract_strategy()
    total_strat = sum(len(v) if isinstance(v, list) else 0 for section in strategy.values() for v in section.values())
    print(f"      ✓ Extracted {total_strat} strategy items")

    # Save files
    print("\n" + "=" * 80)
    print("SAVING JSON FILES")
    print("=" * 80)

    output_files = {
        "suppliers.json": suppliers,
        "consultants.json": consultants,
        "case_studies.json": case_studies,
        "regulatory.json": regulatory,
        "strategy.json": strategy
    }

    for filename, data in output_files.items():
        output_path = DATA_DIR / filename
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

        size = output_path.stat().st_size / 1024
        print(f"✓ {filename:25s} {size:>8.1f} KB")

    # Summary
    print("\n" + "=" * 80)
    print("EXTRACTION SUMMARY")
    print("=" * 80)

    print(f"\n📦 SUPPLIERS ({len(suppliers)} companies)")
    for s in suppliers[:5]:
        print(f"   • {s['name']}")
    if len(suppliers) > 5:
        print(f"   ... and {len(suppliers) - 5} more")

    print(f"\n👥 CONSULTANTS ({len(consultants)} firms)")
    for c in consultants[:5]:
        print(f"   • {c['name']}")
    if len(consultants) > 5:
        print(f"   ... and {len(consultants) - 5} more")

    print(f"\n🏨 CASE STUDIES ({len(case_studies)} hotels)")
    for cs in case_studies[:5]:
        print(f"   • {cs['hotelName']} ({cs['year'] or 'n/a'})")
    if len(case_studies) > 5:
        print(f"   ... and {len(case_studies) - 5} more")

    print(f"\n📋 REGULATORY ({total_reg} items)")
    print(f"   • Fire Safety: {len(regulatory['fireSafety']['bbrRequirements'])} BBR requirements")
    print(f"   • Building Codes: {len(regulatory['buildingCodes']['hotelSpecific'])} hotel-specific")
    print(f"   • Implementation: {len(regulatory['implementation']['tips'])} tips")

    print(f"\n🎯 STRATEGY ({total_strat} items)")
    print(f"   • Market Opportunities: {len(strategy['marketAnalysis']['opportunities'])}")
    print(f"   • Recommendations: {len(strategy['recommendations']['immediate'])} immediate")
    print(f"   • Roadmap: {len(strategy['implementation']['roadmap'])} steps")

    print("\n" + "=" * 80)
    print("✓ EXTRACTION COMPLETE")
    print("=" * 80)
    print(f"\nFiles saved to: {DATA_DIR}/")

if __name__ == "__main__":
    main()
