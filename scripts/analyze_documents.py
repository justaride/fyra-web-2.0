#!/usr/bin/env python3
"""
Document Analysis Script
Extracts and analyzes content from Word documents to understand structure and themes
"""

import os
from docx import Document
from collections import defaultdict
import json

def extract_document_structure(doc_path):
    """Extract headings, paragraphs, and tables from a Word document"""
    try:
        doc = Document(doc_path)
        structure = {
            'filename': os.path.basename(doc_path),
            'headings': [],
            'sections': [],
            'word_count': 0,
            'table_count': len(doc.tables),
            'companies': [],
            'keywords': set()
        }

        current_section = None
        current_content = []
        word_count = 0

        # Extract paragraphs and structure
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            word_count += len(text.split())

            # Check if it's a heading based on style
            style_name = para.style.name.lower()
            if 'heading' in style_name:
                # Save previous section
                if current_section:
                    content_text = ' '.join(current_content)
                    content_words = content_text.split()
                    structure['sections'].append({
                        'heading': current_section,
                        'content_preview': ' '.join(content_words[:50]),
                        'word_count': len(content_words)
                    })

                structure['headings'].append({
                    'level': style_name,
                    'text': text
                })
                current_section = text
                current_content = []
            else:
                current_content.append(text)

        # Add last section
        if current_section:
            content_text = ' '.join(current_content)
            content_words = content_text.split()
            structure['sections'].append({
                'heading': current_section,
                'content_preview': ' '.join(content_words[:50]),
                'word_count': len(content_words)
            })

        structure['word_count'] = word_count

        # Extract table data for first few tables
        structure['tables_preview'] = []
        for i, table in enumerate(doc.tables[:3]):  # First 3 tables
            table_data = []
            for row in table.rows[:5]:  # First 5 rows
                row_data = []
                for cell in row.cells:
                    # Handle merged cells and complex content
                    cell_text = cell.text.strip() if hasattr(cell.text, 'strip') else str(cell.text)
                    row_data.append(cell_text)
                table_data.append(row_data)
            structure['tables_preview'].append({
                'table_number': i + 1,
                'rows_shown': len(table_data),
                'data': table_data
            })

        return structure

    except Exception as e:
        import traceback
        return {
            'filename': os.path.basename(doc_path),
            'error': str(e),
            'traceback': traceback.format_exc()
        }

def analyze_all_documents(directory):
    """Analyze all Word documents in directory"""
    documents = []

    # List of document files in order
    doc_files = [
        "1. COMPREHENSIVE B2B REUSE OPERATORS EXTRACTION.docx",
        "2. FYRA B2B OPERATOR CAPABILITY MATRIX.docx",
        "3. SWEDISH PROJECT MANAGERS & TECHNICAL CONSULTANTS FOR WHOLE-PROJECT HOSPITALITY RENOVATION.docx",
        "4. FRONTRUNNER HOTELS - COMPLETE EXTRACTION.docx",
        "5. PRACTICAL REGULATORY GUIDE_ REUSED MATERIALS IN SWEDISH HOTEL RENOVATIONS.docx",
        "6. STRATEGIC IMPLEMENTATION ROADMAP.docx"
    ]

    for doc_file in doc_files:
        doc_path = os.path.join(directory, doc_file)
        if os.path.exists(doc_path):
            print(f"\nAnalyzing: {doc_file}")
            structure = extract_document_structure(doc_path)
            documents.append(structure)

            # Print summary
            if 'error' not in structure:
                print(f"  - Headings: {len(structure['headings'])}")
                print(f"  - Sections: {len(structure['sections'])}")
                print(f"  - Word count: {structure['word_count']}")
                print(f"  - Tables: {structure['table_count']}")
        else:
            print(f"File not found: {doc_file}")

    return documents

def print_detailed_analysis(documents):
    """Print detailed analysis of all documents"""

    print("\n" + "="*80)
    print("DETAILED DOCUMENT ANALYSIS")
    print("="*80)

    for i, doc in enumerate(documents, 1):
        print(f"\n{'='*80}")
        print(f"DOCUMENT {i}: {doc['filename']}")
        print(f"{'='*80}")

        if 'error' in doc:
            print(f"ERROR: {doc['error']}")
            if 'traceback' in doc:
                print(f"TRACEBACK:\n{doc['traceback']}")
            continue

        print(f"\nSTATISTICS:")
        print(f"  Total Words: {doc['word_count']:,}")
        print(f"  Total Sections: {len(doc['sections'])}")
        print(f"  Total Tables: {doc['table_count']}")

        print(f"\nMAIN SECTIONS ({len(doc['headings'])} headings):")
        for heading in doc['headings']:
            print(f"  {heading['level'].upper()}: {heading['text']}")

        print(f"\nSECTION PREVIEWS:")
        for section in doc['sections'][:10]:  # First 10 sections
            print(f"\n  [{section['heading']}]")
            print(f"    Words: {section['word_count']}")
            print(f"    Preview: {section['content_preview'][:150]}...")

        if doc['tables_preview']:
            print(f"\nTABLE PREVIEWS:")
            for table in doc['tables_preview']:
                print(f"\n  Table {table['table_number']}:")
                for row in table['data'][:3]:  # First 3 rows
                    print(f"    {' | '.join(row[:4])}")  # First 4 columns

    # Cross-document analysis
    print(f"\n{'='*80}")
    print("CROSS-DOCUMENT ANALYSIS")
    print(f"{'='*80}")

    total_words = sum(d.get('word_count', 0) for d in documents if 'error' not in d)
    total_sections = sum(len(d.get('sections', [])) for d in documents if 'error' not in d)
    total_tables = sum(d.get('table_count', 0) for d in documents if 'error' not in d)

    print(f"\nTOTAL STATISTICS:")
    print(f"  Documents: {len(documents)}")
    print(f"  Total words: {total_words:,}")
    print(f"  Total sections: {total_sections}")
    print(f"  Total tables: {total_tables}")

    print(f"\nKEY THEMES IDENTIFIED:")
    all_headings = []
    for doc in documents:
        if 'error' not in doc:
            all_headings.extend([h['text'].lower() for h in doc['headings']])

    # Identify common keywords
    keywords = defaultdict(int)
    for heading in all_headings:
        words = heading.split()
        for word in words:
            if len(word) > 4:  # Only words longer than 4 chars
                keywords[word] += 1

    print("\n  Top keywords from headings:")
    for word, count in sorted(keywords.items(), key=lambda x: x[1], reverse=True)[:20]:
        print(f"    - {word}: {count}")

def suggest_website_structure(documents):
    """Suggest website structure based on document analysis"""

    print(f"\n{'='*80}")
    print("SUGGESTED WEBSITE STRUCTURE")
    print(f"{'='*80}")

    print("""
Based on the 6 documents, here's a suggested website hierarchy:

LEVEL 1: Main Navigation
├── Home
├── About the Project
├── Operators & Services
│   ├── B2B Reuse Operators (Doc 1)
│   ├── Operator Capabilities Matrix (Doc 2)
│   └── Project Managers & Consultants (Doc 3)
├── Case Studies
│   └── Frontrunner Hotels (Doc 4)
├── Resources
│   ├── Regulatory Guide (Doc 5)
│   └── Implementation Roadmap (Doc 6)
└── Contact

LEVEL 2: Detailed Pages

1. OPERATORS & SERVICES SECTION:
   - Overview of circular economy operators
   - Searchable database of operators
   - Capability comparison tool
   - Service categories (by material type)
   - Geographic coverage map

2. CASE STUDIES SECTION:
   - Hotel renovation projects
   - Before/after comparisons
   - Material flows
   - Cost-benefit analysis
   - Lessons learned

3. RESOURCES SECTION:
   - Regulatory compliance guide
   - Step-by-step implementation
   - Templates and tools
   - Best practices
   - FAQ

KEY FEATURES TO IMPLEMENT:
- Search and filter functionality for operators
- Interactive capability matrix
- Downloadable guides and templates
- Case study gallery with metrics
- Resource library
- Contact/inquiry forms
""")

if __name__ == "__main__":
    directory = "/Users/gabrielboen/Downloads/drive-download-20251108T110451Z-1-001"

    print("Starting document analysis...")
    documents = analyze_all_documents(directory)

    print_detailed_analysis(documents)
    suggest_website_structure(documents)

    # Save to JSON for further processing
    output_file = os.path.join(directory, "document_analysis.json")
    with open(output_file, 'w', encoding='utf-8') as f:
        # Convert sets to lists for JSON serialization
        for doc in documents:
            if 'keywords' in doc and isinstance(doc['keywords'], set):
                doc['keywords'] = list(doc['keywords'])
        json.dump(documents, f, indent=2, ensure_ascii=False)

    print(f"\n\nAnalysis saved to: {output_file}")
