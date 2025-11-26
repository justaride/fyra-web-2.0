#!/usr/bin/env python3
"""
REFINED FINAL extraction - properly filters headers vs actual data entries.
"""

import json
import re
from pathlib import Path
from docx import Document

BASE_DIR = Path("/Users/gabrielboen/Downloads/drive-download-20251108T110451Z-1-001")
DATA_DIR = BASE_DIR / "fyra-circular-platform" / "data"

def clean_text(text):
    if not text:
        return ""
    return text.strip().replace('\u200b', '').replace('\n', ' ').replace('  ', ' ').strip()

def extract_table_data(table):
    data = {}
    for row in table.rows:
        if len(row.cells) >= 2:
            key = clean_text(row.cells[0].text)
            value = clean_text(row.cells[1].text)
            if key and value and key != 'Dimension':
                data[key] = value
    return data

def is_actual_company(text):
    """Check if this is an actual company name vs a section header."""
    text_lower = text.lower()

    # Exclude obvious section headers
    exclude_keywords = [
        'tier', 'executive', 'summary', 'immediate action', 'critical finding',
        'gap', 'recommendation', 'partnership model', 'track record',
        'competencies', 'technical scope', 'approach', 'with circular',
        'introduction', 'methodology', 'frontend', 'continued', 'continuation'
    ]

    if any(keyword in text_lower for keyword in exclude_keywords):
        return False

    # Include if it starts with a number (like "1. COMPANY NAME")
    if re.match(r'^\d+\.', text):
        return True

    # Include if it's all caps and medium length (company names often all caps)
    if text.isupper() and 10 < len(text) < 80:
        return True

    # Include if it contains AB, Ltd, AS, Inc (company suffixes)
    if any(suffix in text for suffix in [' AB', ' AS', ' Ltd', ' Inc', ' GmbH', 'Arkitekter', 'Consult']):
        return True

    return False

def extract_suppliers():
    """Extract suppliers."""
    doc_path = BASE_DIR / "1. COMPREHENSIVE B2B REUSE OPERATORS EXTRACTION.docx"
    doc = Document(doc_path)

    suppliers = []
    current_company = None

    for para in doc.paragraphs:
        text = clean_text(para.text)

        if para.style.name == 'Heading 3' and text:
            name = re.sub(r'^\d+\.\s*', '', text)

            if is_actual_company(name):
                if current_company:
                    suppliers.append(current_company)

                current_company = {
                    "id": name.lower().replace(' ', '_').replace('å', 'a').replace('ä', 'a').replace('ö', 'o').replace('/', '_').replace('(', '').replace(')', '').replace('.', '')[:50],
                    "name": name,
                    "description": "",
                    "location": "",
                    "services": [],
                    "capabilities": {"volume": "", "leadTime": "", "inventory": "", "logistics": ""},
                    "contact": {"name": "", "phone": "", "email": "", "website": ""},
                    "certifications": [],
                    "hospitalityReadiness": {"score": "", "strengths": [], "gaps": []},
                    "pricing": "",
                    "projectExamples": []
                }

    if current_company:
        suppliers.append(current_company)

    # Extract from tables
    for table_idx, table in enumerate(doc.tables):
        table_data = extract_table_data(table)
        if not table_data:
            continue

        supplier = None
        for s in suppliers:
            if any(s['name'].split()[0].lower() in str(table_data.get(k, '')).lower()
                   for k in list(table_data.keys())[:5]):
                supplier = s
                break

        if not supplier and table_idx < len(suppliers):
            supplier = suppliers[table_idx]

        if not supplier:
            continue

        for key, value in table_data.items():
            key_lower = key.lower()

            if 'nettside' in key_lower or 'website' in key_lower:
                supplier["contact"]["website"] = value
            elif 'epost' in key_lower or 'email' in key_lower or 'e-post' in key_lower:
                emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', value)
                supplier["contact"]["email"] = emails[0] if emails else value
            elif 'telefon' in key_lower or 'phone' in key_lower:
                supplier["contact"]["phone"] = value
            elif 'ceo' in key_lower or 'kontaktperson' in key_lower or 'founder' in key_lower:
                supplier["contact"]["name"] = value
            elif 'lokasjon' in key_lower or 'location' in key_lower or 'address' in key_lower or 'adresse' in key_lower:
                if not supplier["location"]:
                    supplier["location"] = value
                else:
                    supplier["location"] += "; " + value
            elif 'business model' in key_lower or 'forretningsmodell' in key_lower or 'value proposition' in key_lower:
                supplier["description"] = value
            elif 'services offered' in key_lower or 'tjenester' in key_lower or 'produkter' in key_lower or 'product range' in key_lower:
                if value:
                    supplier["services"] = [s.strip() for s in value.split(',') if s.strip()]
            elif 'volum' in key_lower or 'capacity' in key_lower or 'scale' in key_lower:
                supplier["capabilities"]["volume"] = value
            elif 'lead time' in key_lower or 'leveringstid' in key_lower:
                supplier["capabilities"]["leadTime"] = value
            elif 'inventory' in key_lower or 'lager' in key_lower or 'stock' in key_lower:
                supplier["capabilities"]["inventory"] = value
            elif 'logistics' in key_lower or 'logistikk' in key_lower or 'transport' in key_lower or 'delivery' in key_lower:
                supplier["capabilities"]["logistics"] = value
            elif 'certification' in key_lower or 'sertifisering' in key_lower:
                if value and value.lower() not in ['not documented', 'none', 'n/a']:
                    supplier["certifications"] = [s.strip() for s in value.split(',') if s.strip()]
            elif 'hospitality readiness' in key_lower:
                supplier["hospitalityReadiness"]["score"] = value
            elif 'strengths' in key_lower or 'styrker' in key_lower:
                if value:
                    supplier["hospitalityReadiness"]["strengths"] = [s.strip() for s in value.split('\n') if s.strip()]
            elif 'gaps' in key_lower or 'weaknesses' in key_lower:
                if value:
                    supplier["hospitalityReadiness"]["gaps"] = [s.strip() for s in value.split('\n') if s.strip()]
            elif 'pricing' in key_lower or 'price' in key_lower or 'pris' in key_lower:
                supplier["pricing"] = value
            elif 'project' in key_lower and ('example' in key_lower or 'reference' in key_lower):
                if value and value.lower() not in ['not documented', 'none', 'n/a']:
                    supplier["projectExamples"] = [s.strip() for s in value.split('\n') if s.strip()]

    return suppliers

def extract_consultants():
    """Extract consultants."""
    doc_path = BASE_DIR / "3. SWEDISH PROJECT MANAGERS & TECHNICAL CONSULTANTS FOR WHOLE-PROJECT HOSPITALITY RENOVATION.docx"
    doc = Document(doc_path)

    consultants = []
    current_consultant = None

    for para in doc.paragraphs:
        text = clean_text(para.text)

        if para.style.name == 'Heading 2' and text:
            name = re.sub(r'^\d+\.\s*', '', text)

            if is_actual_company(name):
                if current_consultant:
                    consultants.append(current_consultant)

                current_consultant = {
                    "id": name.lower().replace(' ', '_').replace('å', 'a').replace('ä', 'a').replace('ö', 'o').replace('/', '_').replace('(', '').replace(')', '').replace('.', '').replace('-', '_')[:50],
                    "name": name,
                    "description": "",
                    "services": [],
                    "specializations": [],
                    "hospitalityExperience": {"projects": [], "expertise": ""},
                    "circularEconomyExpertise": "",
                    "contact": {"name": "", "phone": "", "email": "", "website": ""},
                    "certifications": [],
                    "strengths": []
                }

    if current_consultant:
        consultants.append(current_consultant)

    # Extract from tables
    for table_idx, table in enumerate(doc.tables):
        table_data = extract_table_data(table)
        if not table_data:
            continue

        consultant = None
        for c in consultants:
            if any(c['name'].split()[0].lower() in str(table_data.get(k, '')).lower()
                   for k in list(table_data.keys())[:5]):
                consultant = c
                break

        if not consultant and table_idx < len(consultants):
            consultant = consultants[table_idx]

        if not consultant:
            continue

        for key, value in table_data.items():
            key_lower = key.lower()

            if 'website' in key_lower or 'nettside' in key_lower:
                consultant["contact"]["website"] = value
            elif 'email' in key_lower or 'epost' in key_lower:
                emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', value)
                consultant["contact"]["email"] = emails[0] if emails else value
            elif 'phone' in key_lower or 'telefon' in key_lower:
                consultant["contact"]["phone"] = value
            elif 'contact person' in key_lower or 'key contact' in key_lower:
                consultant["contact"]["name"] = value
            elif 'business model' in key_lower or 'company type' in key_lower:
                consultant["description"] = value
            elif 'services' in key_lower or 'capabilities' in key_lower:
                if value:
                    consultant["services"] = [s.strip() for s in value.split(',') if s.strip()][:10]
            elif 'specialization' in key_lower or 'expertise' in key_lower or 'focus' in key_lower:
                if value:
                    consultant["specializations"] = [s.strip() for s in value.split(',') if s.strip()][:10]
            elif 'hospitality' in key_lower and 'project' in key_lower:
                if value and value.lower() not in ['not documented', 'none', 'n/a']:
                    consultant["hospitalityExperience"]["projects"] = [s.strip() for s in value.split('\n') if s.strip()][:10]
            elif 'hospitality' in key_lower and 'experience' in key_lower:
                consultant["hospitalityExperience"]["expertise"] = value
            elif 'circular' in key_lower or 'sustainability' in key_lower:
                if 'certification' in key_lower:
                    consultant["certifications"].append(value)
                else:
                    consultant["circularEconomyExpertise"] = value
            elif 'certification' in key_lower:
                if value and value.lower() not in ['not documented', 'none', 'n/a']:
                    consultant["certifications"].append(value)
            elif 'strength' in key_lower or 'advantage' in key_lower:
                if value:
                    consultant["strengths"] = [s.strip() for s in value.split('\n') if s.strip()][:10]

    return consultants

def is_actual_hotel(text):
    """Check if this is an actual hotel name."""
    text_lower = text.lower()

    # Exclude section headers
    exclude_keywords = [
        'circular economy implementation', 'relevance score', 'sverige', 'denmark',
        'norway', 'netherlands', 'introduction', 'methodology', 'summary',
        'frontrunners', 'continued', 'region', 'stockholm', 'göteborg'
    ]

    if any(keyword in text_lower for keyword in exclude_keywords):
        return False

    # Include if contains hotel keywords
    hotel_keywords = ['hotel', 'hobo', 'hem', 'nobis', 'scandic', 'comfort', 'clarion']
    if any(keyword in text_lower for keyword in hotel_keywords):
        return True

    # Include if it's a proper name (capitalized) and reasonable length
    if text[0].isupper() and 5 < len(text) < 100 and not text.isupper():
        return True

    return False

def extract_case_studies():
    """Extract case studies."""
    doc_path = BASE_DIR / "4. FRONTRUNNER HOTELS - COMPLETE EXTRACTION.docx"
    doc = Document(doc_path)

    case_studies = []
    current_case = None

    for para in doc.paragraphs:
        text = clean_text(para.text)

        if (para.style.name in ['Heading 2', 'Heading 3']) and text:
            if is_actual_hotel(text):
                if current_case:
                    case_studies.append(current_case)

                # Clean hotel name
                hotel_name = text
                # Remove "RELEVANCE SCORE" suffix if present
                hotel_name = re.sub(r'\s*-\s*RELEVANCE SCORE:.*$', '', hotel_name, flags=re.IGNORECASE)

                current_case = {
                    "id": hotel_name.lower().replace(' ', '_').replace('å', 'a').replace('ä', 'a').replace('ö', 'o').replace('-', '_').replace('/', '_')[:50],
                    "hotelName": hotel_name,
                    "location": "",
                    "year": "",
                    "projectSize": "",
                    "category": "",
                    "circularElements": [],
                    "impact": {"co2Savings": "", "circularPercentage": "", "costSavings": ""},
                    "architects": [],
                    "partners": [],
                    "challenges": [],
                    "outcomes": [],
                    "relevance": ""
                }

    if current_case:
        case_studies.append(current_case)

    # Extract from tables
    for table in doc.tables:
        table_data = extract_table_data(table)
        if not table_data:
            continue

        case_study = None
        for cs in case_studies:
            if any(cs['hotelName'].split()[0].lower() in str(table_data.get(k, '')).lower()
                   for k in list(table_data.keys())[:5]):
                case_study = cs
                break

        if not case_study and case_studies:
            case_study = case_studies[-1]

        if not case_study:
            continue

        for key, value in table_data.items():
            key_lower = key.lower()

            if 'location' in key_lower or 'city' in key_lower:
                case_study["location"] = value
            elif 'year' in key_lower or 'opened' in key_lower or 'completed' in key_lower:
                year_match = re.search(r'\b(19|20)\d{2}\b', value)
                case_study["year"] = year_match.group() if year_match else value
            elif 'size' in key_lower or 'scale' in key_lower or 'rooms' in key_lower:
                case_study["projectSize"] = value
            elif 'category' in key_lower or 'classification' in key_lower or 'type' in key_lower:
                case_study["category"] = value
            elif 'circular' in key_lower or 'reused' in key_lower or 'reuse' in key_lower or 'upcycled' in key_lower or 'material' in key_lower:
                if value and value.lower() not in ['not documented', 'none', 'n/a']:
                    case_study["circularElements"].append(f"{key}: {value}")
            elif 'co2' in key_lower or 'carbon' in key_lower or 'emission' in key_lower:
                case_study["impact"]["co2Savings"] = value
            elif '%' in value and ('circular' in key_lower or 'recycl' in key_lower or 'reuse' in key_lower):
                case_study["impact"]["circularPercentage"] = value
            elif 'cost' in key_lower and ('saving' in key_lower or 'reduction' in key_lower):
                case_study["impact"]["costSavings"] = value
            elif 'architect' in key_lower or 'designer' in key_lower:
                if value and value.lower() not in ['not documented', 'none', 'n/a']:
                    case_study["architects"].append(value)
            elif 'supplier' in key_lower or 'contractor' in key_lower or 'partner' in key_lower or 'consultant' in key_lower:
                if value and value.lower() not in ['not documented', 'none', 'n/a']:
                    case_study["partners"].append(value)
            elif 'challenge' in key_lower:
                if value:
                    case_study["challenges"].append(value)
            elif 'outcome' in key_lower or 'result' in key_lower:
                if value:
                    case_study["outcomes"].append(value)
            elif 'relevance' in key_lower or 'lesson' in key_lower or 'justification' in key_lower:
                case_study["relevance"] = value

    return case_studies

def extract_regulatory():
    """Extract regulatory."""
    doc_path = BASE_DIR / "5. PRACTICAL REGULATORY GUIDE_ REUSED MATERIALS IN SWEDISH HOTEL RENOVATIONS.docx"
    doc = Document(doc_path)

    regulatory = {
        "fireSafety": {"overview": "", "bbrRequirements": [], "testingStandards": [], "documentation": []},
        "buildingCodes": {"overview": "", "hotelSpecific": [], "materialRequirements": []},
        "bvbStandards": {"overview": "", "scoring": [], "certificationProcess": []},
        "documentation": {"required": [], "recommended": [], "templates": []},
        "implementation": {"tips": [], "bestPractices": [], "commonChallenges": []}
    }

    for table in doc.tables:
        table_data = extract_table_data(table)
        for key, value in table_data.items():
            key_lower = key.lower()

            if 'fire' in key_lower or 'brand' in key_lower:
                if 'bbr' in value.lower():
                    regulatory["fireSafety"]["bbrRequirements"].append(f"{key}: {value}")
                elif 'test' in value.lower():
                    regulatory["fireSafety"]["testingStandards"].append(f"{key}: {value}")
                else:
                    regulatory["fireSafety"]["documentation"].append(f"{key}: {value}")
            elif 'bbr' in key_lower or 'building code' in key_lower:
                if 'hotel' in value.lower():
                    regulatory["buildingCodes"]["hotelSpecific"].append(f"{key}: {value}")
                else:
                    regulatory["buildingCodes"]["materialRequirements"].append(f"{key}: {value}")
            elif 'bvb' in key_lower:
                if 'point' in value.lower() or 'score' in value.lower():
                    regulatory["bvbStandards"]["scoring"].append(f"{key}: {value}")
                else:
                    regulatory["bvbStandards"]["certificationProcess"].append(f"{key}: {value}")
            elif 'documentation' in key_lower:
                if 'required' in value.lower():
                    regulatory["documentation"]["required"].append(f"{key}: {value}")
                elif 'recommended' in value.lower():
                    regulatory["documentation"]["recommended"].append(f"{key}: {value}")
                else:
                    regulatory["documentation"]["templates"].append(f"{key}: {value}")
            elif 'tip' in key_lower or 'practice' in key_lower or 'challenge' in key_lower:
                if 'challenge' in key_lower:
                    regulatory["implementation"]["commonChallenges"].append(f"{key}: {value}")
                elif 'best practice' in key_lower:
                    regulatory["implementation"]["bestPractices"].append(f"{key}: {value}")
                else:
                    regulatory["implementation"]["tips"].append(f"{key}: {value}")

    return regulatory

def extract_strategy():
    """Extract strategy."""
    doc_path = BASE_DIR / "6. STRATEGIC IMPLEMENTATION ROADMAP.docx"
    doc = Document(doc_path)

    strategy = {
        "marketAnalysis": {"swedenOverview": "", "hotelSector": "", "circularEconomyTrends": "", "opportunities": [], "challenges": []},
        "fyraPositioning": {"valueProposition": "", "differentiators": [], "targetMarket": "", "competitiveAdvantages": []},
        "recommendations": {"immediate": [], "shortTerm": [], "longTerm": []},
        "implementation": {"roadmap": [], "milestones": [], "resources": [], "successMetrics": []}
    }

    for table in doc.tables:
        table_data = extract_table_data(table)
        for key, value in table_data.items():
            key_lower = key.lower()

            if 'market' in key_lower or 'sweden' in key_lower or 'nordic' in key_lower:
                if 'opportunit' in key_lower:
                    strategy["marketAnalysis"]["opportunities"].append(f"{key}: {value}")
                elif 'challenge' in key_lower:
                    strategy["marketAnalysis"]["challenges"].append(f"{key}: {value}")
                elif 'hotel' in key_lower and not strategy["marketAnalysis"]["hotelSector"]:
                    strategy["marketAnalysis"]["hotelSector"] = value
                elif 'circular' in key_lower and not strategy["marketAnalysis"]["circularEconomyTrends"]:
                    strategy["marketAnalysis"]["circularEconomyTrends"] = value
            elif 'fyra' in key_lower or 'position' in key_lower or 'value proposition' in key_lower:
                if 'different' in key_lower or 'unique' in key_lower:
                    strategy["fyraPositioning"]["differentiators"].append(f"{key}: {value}")
                elif 'advantage' in key_lower or 'competitive' in key_lower:
                    strategy["fyraPositioning"]["competitiveAdvantages"].append(f"{key}: {value}")
                elif 'target' in key_lower and not strategy["fyraPositioning"]["targetMarket"]:
                    strategy["fyraPositioning"]["targetMarket"] = value
            elif 'recommend' in key_lower or 'action' in key_lower or 'priority' in key_lower:
                if 'immediate' in key_lower or 'short' in key_lower or 'month 1' in key_lower:
                    strategy["recommendations"]["immediate"].append(f"{key}: {value}")
                elif 'medium' in key_lower or 'month 3' in key_lower:
                    strategy["recommendations"]["shortTerm"].append(f"{key}: {value}")
                elif 'long' in key_lower or 'year' in key_lower:
                    strategy["recommendations"]["longTerm"].append(f"{key}: {value}")
                else:
                    strategy["recommendations"]["immediate"].append(f"{key}: {value}")
            elif 'roadmap' in key_lower or 'timeline' in key_lower or 'phase' in key_lower or 'step' in key_lower:
                strategy["implementation"]["roadmap"].append(f"{key}: {value}")
            elif 'milestone' in key_lower:
                strategy["implementation"]["milestones"].append(f"{key}: {value}")
            elif 'resource' in key_lower:
                strategy["implementation"]["resources"].append(f"{key}: {value}")
            elif 'metric' in key_lower or 'kpi' in key_lower:
                strategy["implementation"]["successMetrics"].append(f"{key}: {value}")

    return strategy

def main():
    print("=" * 80)
    print("FYRA CIRCULAR PLATFORM - REFINED FINAL EXTRACTION")
    print("=" * 80)

    suppliers = extract_suppliers()
    consultants = extract_consultants()
    case_studies = extract_case_studies()
    regulatory = extract_regulatory()
    strategy = extract_strategy()

    print(f"\n✓ Extracted {len(suppliers)} suppliers")
    print(f"✓ Extracted {len(consultants)} consultants")
    print(f"✓ Extracted {len(case_studies)} case studies")

    output_files = {
        "suppliers.json": suppliers,
        "consultants.json": consultants,
        "case_studies.json": case_studies,
        "regulatory.json": regulatory,
        "strategy.json": strategy
    }

    print("\n" + "=" * 80)
    print("SAVING FILES")
    print("=" * 80)

    for filename, data in output_files.items():
        output_path = DATA_DIR / filename
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        size = output_path.stat().st_size / 1024
        print(f"✓ {filename:25s} {size:>8.1f} KB")

    print("\n" + "=" * 80)
    print("SUMMARY")
    print("=" * 80)

    print(f"\n📦 SUPPLIERS ({len(suppliers)})")
    for s in suppliers[:5]:
        email = "✓" if s['contact']['email'] else "✗"
        certs = len(s['certifications'])
        print(f"   {s['name'][:45]:45s} | Email: {email} | Certs: {certs}")
    if len(suppliers) > 5:
        print(f"   ... +{len(suppliers) - 5} more")

    print(f"\n👥 CONSULTANTS ({len(consultants)})")
    for c in consultants[:5]:
        email = "✓" if c['contact']['email'] or c['contact']['website'] else "✗"
        hosp = len(c['hospitalityExperience']['projects'])
        print(f"   {c['name'][:45]:45s} | Contact: {email} | Hosp: {hosp}")
    if len(consultants) > 5:
        print(f"   ... +{len(consultants) - 5} more")

    print(f"\n🏨 CASE STUDIES ({len(case_studies)})")
    for cs in case_studies[:5]:
        year = cs['year'][:4] if cs['year'] else 'n/a '
        circ = len(cs['circularElements'])
        print(f"   {cs['hotelName'][:45]:45s} | {year} | Circular: {circ}")
    if len(case_studies) > 5:
        print(f"   ... +{len(case_studies) - 5} more")

    print("\n" + "=" * 80)
    print("✓ COMPLETE")
    print("=" * 80)

if __name__ == "__main__":
    main()
