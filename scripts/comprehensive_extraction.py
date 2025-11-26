#!/usr/bin/env python3
"""
Comprehensive extraction of all data from Fyra Circular Platform documents.
Extracts suppliers, consultants, case studies, regulatory info, and strategy.
"""

import json
import re
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# Paths
BASE_DIR = Path("/Users/gabrielboen/Downloads/drive-download-20251108T110451Z-1-001")
DATA_DIR = BASE_DIR / "fyra-circular-platform" / "data"

# Document paths
DOCS = {
    "suppliers": BASE_DIR / "1. COMPREHENSIVE B2B REUSE OPERATORS EXTRACTION.docx",
    "matrix": BASE_DIR / "2. FYRA B2B OPERATOR CAPABILITY MATRIX.docx",
    "consultants": BASE_DIR / "3. SWEDISH PROJECT MANAGERS & TECHNICAL CONSULTANTS FOR WHOLE-PROJECT HOSPITALITY RENOVATION.docx",
    "hotels": BASE_DIR / "4. FRONTRUNNER HOTELS - COMPLETE EXTRACTION.docx",
    "regulatory": BASE_DIR / "5. PRACTICAL REGULATORY GUIDE_ REUSED MATERIALS IN SWEDISH HOTEL RENOVATIONS.docx",
    "strategy": BASE_DIR / "6. STRATEGIC IMPLEMENTATION ROADMAP.docx"
}

def clean_text(text):
    """Clean and normalize text."""
    if not text:
        return ""
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\u200b', '', text)  # Remove zero-width spaces
    return text

def extract_contact_info(text):
    """Extract contact information from text."""
    contacts = {
        "emails": [],
        "phones": [],
        "websites": []
    }

    # Extract emails
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    contacts["emails"] = list(set(re.findall(email_pattern, text)))

    # Extract phone numbers (Swedish format)
    phone_pattern = r'\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}'
    phones = re.findall(phone_pattern, text)
    contacts["phones"] = [clean_text(p) for p in phones if len(p.replace('-', '').replace(' ', '').replace('.', '')) >= 8]

    # Extract websites
    website_pattern = r'(?:https?://)?(?:www\.)?[a-zA-Z0-9-]+\.[a-zA-Z]{2,}(?:/[^\s]*)?'
    websites = re.findall(website_pattern, text)
    contacts["websites"] = [w for w in websites if '@' not in w]

    return contacts

def get_all_text(doc):
    """Extract all text from document including tables."""
    full_text = []

    for element in doc.element.body:
        if isinstance(element, Paragraph):
            full_text.append(element.text)
        elif hasattr(element, 'tag') and 'tbl' in element.tag:
            # This is a table
            for row in element.iter():
                if hasattr(row, 'text'):
                    full_text.append(row.text)

    return '\n'.join(full_text)

def extract_suppliers(doc_path):
    """Extract B2B suppliers from Document 1."""
    doc = Document(doc_path)
    suppliers = []
    current_supplier = None

    print(f"\nExtracting suppliers from: {doc_path.name}")

    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"

    # Split by company entries (look for patterns like company names in headers)
    sections = []
    current_section = {"title": "", "content": [], "tables": []}

    for block in doc.element.body:
        if hasattr(block, 'tag'):
            if 'p' in block.tag:
                para = Paragraph(block, doc)
                text = clean_text(para.text)
                if text:
                    # Check if this is a header (bold, larger font, or heading style)
                    is_header = False
                    if para.style.name.startswith('Heading') or (para.runs and para.runs[0].bold and len(text) < 100):
                        is_header = True

                    if is_header and current_section["content"]:
                        sections.append(current_section)
                        current_section = {"title": text, "content": [], "tables": []}
                    elif is_header:
                        current_section["title"] = text
                    else:
                        current_section["content"].append(text)

            elif 'tbl' in block.tag:
                table = Table(block, doc)
                table_data = []
                for row in table.rows:
                    row_data = [clean_text(cell.text) for cell in row.cells]
                    if any(row_data):
                        table_data.append(row_data)
                if table_data:
                    current_section["tables"].append(table_data)

    if current_section["content"] or current_section["tables"]:
        sections.append(current_section)

    # Process sections into suppliers
    for section in sections:
        content_text = "\n".join(section["content"])

        # Skip non-supplier sections
        if not section["title"] or any(skip in section["title"].lower() for skip in
            ['introduction', 'overview', 'summary', 'table of contents', 'methodology']):
            continue

        # Extract company name from title
        company_name = section["title"]

        # Clean up common prefixes
        for prefix in ['Company:', 'Operator:', 'Profile:']:
            if company_name.startswith(prefix):
                company_name = company_name[len(prefix):].strip()

        if not company_name or len(company_name) > 100:
            continue

        contacts = extract_contact_info(content_text + "\n" + str(section["tables"]))

        supplier = {
            "id": company_name.lower().replace(' ', '_').replace('å', 'a').replace('ä', 'a').replace('ö', 'o'),
            "name": company_name,
            "description": "",
            "location": "",
            "services": [],
            "capabilities": {
                "volume": "",
                "leadTime": "",
                "inventory": "",
                "logistics": ""
            },
            "contact": {
                "name": "",
                "phone": contacts["phones"][0] if contacts["phones"] else "",
                "email": contacts["emails"][0] if contacts["emails"] else "",
                "website": contacts["websites"][0] if contacts["websites"] else ""
            },
            "certifications": [],
            "hospitalityReadiness": {
                "score": "",
                "strengths": [],
                "gaps": []
            },
            "pricing": "",
            "projectExamples": []
        }

        # Parse content for specific fields
        lines = content_text.split('\n')
        current_field = None

        for line in lines:
            line = clean_text(line)
            if not line:
                continue

            lower_line = line.lower()

            # Location
            if 'location' in lower_line or 'address' in lower_line or 'based in' in lower_line:
                if ':' in line:
                    supplier["location"] = clean_text(line.split(':', 1)[1])
                else:
                    supplier["location"] = line

            # Description
            elif 'description' in lower_line or 'overview' in lower_line or 'about' in lower_line:
                if ':' in line:
                    supplier["description"] = clean_text(line.split(':', 1)[1])
                    current_field = "description"

            # Services
            elif 'service' in lower_line or 'offer' in lower_line or 'product' in lower_line:
                if ':' in line:
                    services_text = line.split(':', 1)[1]
                    supplier["services"] = [s.strip() for s in services_text.split(',') if s.strip()]
                current_field = "services"

            # Capabilities
            elif 'volume' in lower_line or 'capacity' in lower_line:
                if ':' in line:
                    supplier["capabilities"]["volume"] = clean_text(line.split(':', 1)[1])

            elif 'lead time' in lower_line or 'leadtime' in lower_line:
                if ':' in line:
                    supplier["capabilities"]["leadTime"] = clean_text(line.split(':', 1)[1])

            elif 'inventory' in lower_line or 'stock' in lower_line:
                if ':' in line:
                    supplier["capabilities"]["inventory"] = clean_text(line.split(':', 1)[1])

            elif 'logistics' in lower_line or 'delivery' in lower_line or 'transport' in lower_line:
                if ':' in line:
                    supplier["capabilities"]["logistics"] = clean_text(line.split(':', 1)[1])

            # Certifications
            elif 'certification' in lower_line or 'certified' in lower_line or 'accreditation' in lower_line:
                if ':' in line:
                    certs = line.split(':', 1)[1]
                    supplier["certifications"] = [c.strip() for c in certs.split(',') if c.strip()]
                current_field = "certifications"

            # Hospitality readiness
            elif 'hospitality' in lower_line and ('readiness' in lower_line or 'score' in lower_line):
                if ':' in line:
                    supplier["hospitalityReadiness"]["score"] = clean_text(line.split(':', 1)[1])

            elif 'strength' in lower_line:
                current_field = "strengths"

            elif 'gap' in lower_line or 'weakness' in lower_line or 'challenge' in lower_line:
                current_field = "gaps"

            # Pricing
            elif 'pric' in lower_line or 'cost' in lower_line or 'rate' in lower_line:
                if ':' in line:
                    supplier["pricing"] = clean_text(line.split(':', 1)[1])

            # Project examples
            elif 'project' in lower_line and ('example' in lower_line or 'case' in lower_line):
                current_field = "projects"

            # Continue adding to current field
            elif current_field:
                if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                    item = line[1:].strip()
                    if current_field == "services":
                        supplier["services"].append(item)
                    elif current_field == "certifications":
                        supplier["certifications"].append(item)
                    elif current_field == "strengths":
                        supplier["hospitalityReadiness"]["strengths"].append(item)
                    elif current_field == "gaps":
                        supplier["hospitalityReadiness"]["gaps"].append(item)
                    elif current_field == "projects":
                        supplier["projectExamples"].append(item)
                elif current_field == "description" and not any(keyword in lower_line for keyword in
                    ['service', 'location', 'contact', 'certification', 'capability']):
                    supplier["description"] += " " + line

        # Process tables
        for table in section["tables"]:
            for row in table:
                if len(row) >= 2:
                    key = clean_text(row[0]).lower()
                    value = clean_text(row[1])

                    if 'service' in key or 'product' in key:
                        if value and value not in supplier["services"]:
                            supplier["services"].append(value)
                    elif 'certification' in key:
                        if value and value not in supplier["certifications"]:
                            supplier["certifications"].append(value)
                    elif 'location' in key or 'address' in key:
                        supplier["location"] = value

        # Clean up
        supplier["description"] = clean_text(supplier["description"])
        if len(supplier["description"]) > 500:
            supplier["description"] = supplier["description"][:500] + "..."

        suppliers.append(supplier)
        print(f"  - Extracted: {supplier['name']}")

    return suppliers

def extract_consultants(doc_path):
    """Extract consultants/PMs from Document 3."""
    doc = Document(doc_path)
    consultants = []

    print(f"\nExtracting consultants from: {doc_path.name}")

    sections = []
    current_section = {"title": "", "content": [], "tables": []}

    for block in doc.element.body:
        if hasattr(block, 'tag'):
            if 'p' in block.tag:
                para = Paragraph(block, doc)
                text = clean_text(para.text)
                if text:
                    is_header = para.style.name.startswith('Heading') or (para.runs and para.runs[0].bold and len(text) < 100)

                    if is_header and current_section["content"]:
                        sections.append(current_section)
                        current_section = {"title": text, "content": [], "tables": []}
                    elif is_header:
                        current_section["title"] = text
                    else:
                        current_section["content"].append(text)

            elif 'tbl' in block.tag:
                table = Table(block, doc)
                table_data = []
                for row in table.rows:
                    row_data = [clean_text(cell.text) for cell in row.cells]
                    if any(row_data):
                        table_data.append(row_data)
                if table_data:
                    current_section["tables"].append(table_data)

    if current_section["content"] or current_section["tables"]:
        sections.append(current_section)

    for section in sections:
        content_text = "\n".join(section["content"])

        if not section["title"] or any(skip in section["title"].lower() for skip in
            ['introduction', 'overview', 'summary', 'table of contents']):
            continue

        company_name = section["title"]
        for prefix in ['Company:', 'Firm:', 'Consultant:']:
            if company_name.startswith(prefix):
                company_name = company_name[len(prefix):].strip()

        if not company_name or len(company_name) > 100:
            continue

        contacts = extract_contact_info(content_text + "\n" + str(section["tables"]))

        consultant = {
            "id": company_name.lower().replace(' ', '_').replace('å', 'a').replace('ä', 'a').replace('ö', 'o'),
            "name": company_name,
            "description": "",
            "services": [],
            "specializations": [],
            "hospitalityExperience": {
                "projects": [],
                "expertise": ""
            },
            "circularEconomyExpertise": "",
            "contact": {
                "name": "",
                "phone": contacts["phones"][0] if contacts["phones"] else "",
                "email": contacts["emails"][0] if contacts["emails"] else "",
                "website": contacts["websites"][0] if contacts["websites"] else ""
            },
            "certifications": [],
            "strengths": []
        }

        lines = content_text.split('\n')
        current_field = None

        for line in lines:
            line = clean_text(line)
            if not line:
                continue

            lower_line = line.lower()

            if 'description' in lower_line or 'overview' in lower_line or 'about' in lower_line:
                if ':' in line:
                    consultant["description"] = clean_text(line.split(':', 1)[1])
                    current_field = "description"

            elif 'service' in lower_line or 'offer' in lower_line:
                if ':' in line:
                    services = line.split(':', 1)[1]
                    consultant["services"] = [s.strip() for s in services.split(',') if s.strip()]
                current_field = "services"

            elif 'specialization' in lower_line or 'expertise' in lower_line:
                if ':' in line:
                    specs = line.split(':', 1)[1]
                    consultant["specializations"] = [s.strip() for s in specs.split(',') if s.strip()]
                current_field = "specializations"

            elif 'hospitality' in lower_line and ('project' in lower_line or 'experience' in lower_line):
                current_field = "hospitality_projects"

            elif 'circular' in lower_line and 'econom' in lower_line:
                if ':' in line:
                    consultant["circularEconomyExpertise"] = clean_text(line.split(':', 1)[1])
                current_field = "circular"

            elif 'certification' in lower_line or 'certified' in lower_line:
                if ':' in line:
                    certs = line.split(':', 1)[1]
                    consultant["certifications"] = [c.strip() for c in certs.split(',') if c.strip()]
                current_field = "certifications"

            elif 'strength' in lower_line:
                current_field = "strengths"

            elif current_field:
                if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                    item = line[1:].strip()
                    if current_field == "services":
                        consultant["services"].append(item)
                    elif current_field == "specializations":
                        consultant["specializations"].append(item)
                    elif current_field == "hospitality_projects":
                        consultant["hospitalityExperience"]["projects"].append(item)
                    elif current_field == "certifications":
                        consultant["certifications"].append(item)
                    elif current_field == "strengths":
                        consultant["strengths"].append(item)
                elif current_field == "description":
                    consultant["description"] += " " + line
                elif current_field == "circular":
                    consultant["circularEconomyExpertise"] += " " + line

        consultant["description"] = clean_text(consultant["description"])
        consultant["circularEconomyExpertise"] = clean_text(consultant["circularEconomyExpertise"])

        consultants.append(consultant)
        print(f"  - Extracted: {consultant['name']}")

    return consultants

def extract_case_studies(doc_path):
    """Extract hotel case studies from Document 4."""
    doc = Document(doc_path)
    case_studies = []

    print(f"\nExtracting case studies from: {doc_path.name}")

    sections = []
    current_section = {"title": "", "content": [], "tables": []}

    for block in doc.element.body:
        if hasattr(block, 'tag'):
            if 'p' in block.tag:
                para = Paragraph(block, doc)
                text = clean_text(para.text)
                if text:
                    is_header = para.style.name.startswith('Heading') or (para.runs and para.runs[0].bold and len(text) < 150)

                    if is_header and current_section["content"]:
                        sections.append(current_section)
                        current_section = {"title": text, "content": [], "tables": []}
                    elif is_header:
                        current_section["title"] = text
                    else:
                        current_section["content"].append(text)

            elif 'tbl' in block.tag:
                table = Table(block, doc)
                table_data = []
                for row in table.rows:
                    row_data = [clean_text(cell.text) for cell in row.cells]
                    if any(row_data):
                        table_data.append(row_data)
                if table_data:
                    current_section["tables"].append(table_data)

    if current_section["content"] or current_section["tables"]:
        sections.append(current_section)

    for section in sections:
        content_text = "\n".join(section["content"])

        if not section["title"] or any(skip in section["title"].lower() for skip in
            ['introduction', 'overview', 'summary', 'table of contents', 'methodology']):
            continue

        # Hotel name from title
        hotel_name = section["title"]
        for prefix in ['Hotel:', 'Case Study:', 'Project:']:
            if hotel_name.startswith(prefix):
                hotel_name = hotel_name[len(prefix):].strip()

        if not hotel_name or len(hotel_name) > 150:
            continue

        case_study = {
            "id": hotel_name.lower().replace(' ', '_').replace('å', 'a').replace('ä', 'a').replace('ö', 'o')[:50],
            "hotelName": hotel_name,
            "location": "",
            "year": "",
            "projectSize": "",
            "category": "",
            "circularElements": [],
            "impact": {
                "co2Savings": "",
                "circularPercentage": "",
                "costSavings": ""
            },
            "architects": [],
            "partners": [],
            "challenges": [],
            "outcomes": [],
            "relevance": ""
        }

        lines = content_text.split('\n')
        current_field = None

        for line in lines:
            line = clean_text(line)
            if not line:
                continue

            lower_line = line.lower()

            if 'location' in lower_line or 'city' in lower_line or 'country' in lower_line:
                if ':' in line:
                    case_study["location"] = clean_text(line.split(':', 1)[1])

            elif 'year' in lower_line or 'completed' in lower_line or 'opened' in lower_line:
                if ':' in line:
                    case_study["year"] = clean_text(line.split(':', 1)[1])
                # Extract year from text
                year_match = re.search(r'\b(19|20)\d{2}\b', line)
                if year_match and not case_study["year"]:
                    case_study["year"] = year_match.group()

            elif 'size' in lower_line or 'scale' in lower_line or 'm2' in lower_line or 'sqm' in lower_line:
                if ':' in line:
                    case_study["projectSize"] = clean_text(line.split(':', 1)[1])

            elif 'category' in lower_line or 'type' in lower_line or 'classification' in lower_line:
                if ':' in line:
                    case_study["category"] = clean_text(line.split(':', 1)[1])

            elif 'circular element' in lower_line or 'reused material' in lower_line or 'reuse' in lower_line:
                current_field = "circular_elements"

            elif 'co2' in lower_line or 'carbon' in lower_line or 'emission' in lower_line:
                if ':' in line:
                    case_study["impact"]["co2Savings"] = clean_text(line.split(':', 1)[1])

            elif 'circular' in lower_line and '%' in line:
                if ':' in line:
                    case_study["impact"]["circularPercentage"] = clean_text(line.split(':', 1)[1])

            elif 'cost' in lower_line and ('saving' in lower_line or 'reduction' in lower_line):
                if ':' in line:
                    case_study["impact"]["costSavings"] = clean_text(line.split(':', 1)[1])

            elif 'architect' in lower_line:
                if ':' in line:
                    archs = line.split(':', 1)[1]
                    case_study["architects"] = [a.strip() for a in archs.split(',') if a.strip()]
                current_field = "architects"

            elif 'partner' in lower_line or 'collaborator' in lower_line:
                current_field = "partners"

            elif 'challenge' in lower_line or 'difficulty' in lower_line or 'obstacle' in lower_line:
                current_field = "challenges"

            elif 'outcome' in lower_line or 'result' in lower_line or 'achievement' in lower_line:
                current_field = "outcomes"

            elif 'relevance' in lower_line or 'lesson' in lower_line:
                if ':' in line:
                    case_study["relevance"] = clean_text(line.split(':', 1)[1])
                current_field = "relevance"

            elif current_field:
                if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                    item = line[1:].strip()
                    if current_field == "circular_elements":
                        case_study["circularElements"].append(item)
                    elif current_field == "architects":
                        case_study["architects"].append(item)
                    elif current_field == "partners":
                        case_study["partners"].append(item)
                    elif current_field == "challenges":
                        case_study["challenges"].append(item)
                    elif current_field == "outcomes":
                        case_study["outcomes"].append(item)
                elif current_field == "relevance":
                    case_study["relevance"] += " " + line

        case_study["relevance"] = clean_text(case_study["relevance"])

        case_studies.append(case_study)
        print(f"  - Extracted: {case_study['hotelName']}")

    return case_studies

def extract_regulatory(doc_path):
    """Extract regulatory information from Document 5."""
    doc = Document(doc_path)

    print(f"\nExtracting regulatory info from: {doc_path.name}")

    regulatory = {
        "fireSafety": {
            "overview": "",
            "bbrRequirements": [],
            "testingStandards": [],
            "documentation": []
        },
        "buildingCodes": {
            "overview": "",
            "hotelSpecific": [],
            "materialRequirements": []
        },
        "bvbStandards": {
            "overview": "",
            "scoring": [],
            "certificationProcess": []
        },
        "documentation": {
            "required": [],
            "recommended": [],
            "templates": []
        },
        "implementation": {
            "tips": [],
            "bestPractices": [],
            "commonChallenges": []
        }
    }

    current_section = None
    current_subsection = None

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        lower_text = text.lower()

        # Identify main sections
        if para.style.name.startswith('Heading 1'):
            if 'fire' in lower_text and 'safety' in lower_text:
                current_section = "fireSafety"
            elif 'building code' in lower_text or 'bbr' in lower_text:
                current_section = "buildingCodes"
            elif 'bvb' in lower_text or 'environmental' in lower_text:
                current_section = "bvbStandards"
            elif 'documentation' in lower_text:
                current_section = "documentation"
            elif 'implementation' in lower_text or 'practical' in lower_text:
                current_section = "implementation"

        # Identify subsections
        elif para.style.name.startswith('Heading 2') and current_section:
            if 'overview' in lower_text or 'introduction' in lower_text:
                current_subsection = "overview"
            elif 'requirement' in lower_text:
                current_subsection = "requirements"
            elif 'testing' in lower_text or 'standard' in lower_text:
                current_subsection = "testing"
            elif 'documentation' in lower_text:
                current_subsection = "documentation"
            elif 'tip' in lower_text or 'advice' in lower_text:
                current_subsection = "tips"
            elif 'practice' in lower_text:
                current_subsection = "practices"
            elif 'challenge' in lower_text:
                current_subsection = "challenges"
            elif 'scoring' in lower_text or 'point' in lower_text:
                current_subsection = "scoring"

        # Add content
        elif current_section and text:
            if text.startswith('•') or text.startswith('-') or text.startswith('*'):
                item = text[1:].strip()

                if current_section == "fireSafety":
                    if current_subsection == "requirements" or 'bbr' in lower_text:
                        regulatory["fireSafety"]["bbrRequirements"].append(item)
                    elif current_subsection == "testing":
                        regulatory["fireSafety"]["testingStandards"].append(item)
                    elif current_subsection == "documentation":
                        regulatory["fireSafety"]["documentation"].append(item)

                elif current_section == "buildingCodes":
                    if 'hotel' in lower_text or current_subsection == "requirements":
                        regulatory["buildingCodes"]["hotelSpecific"].append(item)
                    else:
                        regulatory["buildingCodes"]["materialRequirements"].append(item)

                elif current_section == "bvbStandards":
                    if current_subsection == "scoring":
                        regulatory["bvbStandards"]["scoring"].append(item)
                    else:
                        regulatory["bvbStandards"]["certificationProcess"].append(item)

                elif current_section == "documentation":
                    if 'required' in lower_text:
                        regulatory["documentation"]["required"].append(item)
                    elif 'recommended' in lower_text:
                        regulatory["documentation"]["recommended"].append(item)
                    else:
                        regulatory["documentation"]["templates"].append(item)

                elif current_section == "implementation":
                    if current_subsection == "tips":
                        regulatory["implementation"]["tips"].append(item)
                    elif current_subsection == "practices":
                        regulatory["implementation"]["bestPractices"].append(item)
                    elif current_subsection == "challenges":
                        regulatory["implementation"]["commonChallenges"].append(item)

            elif current_subsection == "overview":
                if current_section == "fireSafety":
                    regulatory["fireSafety"]["overview"] += " " + text
                elif current_section == "buildingCodes":
                    regulatory["buildingCodes"]["overview"] += " " + text
                elif current_section == "bvbStandards":
                    regulatory["bvbStandards"]["overview"] += " " + text

    # Clean up overviews
    regulatory["fireSafety"]["overview"] = clean_text(regulatory["fireSafety"]["overview"])
    regulatory["buildingCodes"]["overview"] = clean_text(regulatory["buildingCodes"]["overview"])
    regulatory["bvbStandards"]["overview"] = clean_text(regulatory["bvbStandards"]["overview"])

    print(f"  - Extracted {len(regulatory['fireSafety']['bbrRequirements'])} fire safety requirements")
    print(f"  - Extracted {len(regulatory['buildingCodes']['hotelSpecific'])} building code items")
    print(f"  - Extracted {len(regulatory['implementation']['tips'])} implementation tips")

    return regulatory

def extract_strategy(doc_path):
    """Extract strategic information from Document 6."""
    doc = Document(doc_path)

    print(f"\nExtracting strategy from: {doc_path.name}")

    strategy = {
        "marketAnalysis": {
            "swedenOverview": "",
            "hotelSector": "",
            "circularEconomyTrends": "",
            "opportunities": [],
            "challenges": []
        },
        "fyraPositioning": {
            "valueProposition": "",
            "differentiators": [],
            "targetMarket": "",
            "competitiveAdvantages": []
        },
        "recommendations": {
            "immediate": [],
            "shortTerm": [],
            "longTerm": []
        },
        "implementation": {
            "roadmap": [],
            "milestones": [],
            "resources": [],
            "successMetrics": []
        }
    }

    current_section = None
    current_subsection = None

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        lower_text = text.lower()

        # Identify main sections
        if para.style.name.startswith('Heading 1'):
            if 'market' in lower_text and 'analysis' in lower_text:
                current_section = "marketAnalysis"
            elif 'fyra' in lower_text or 'position' in lower_text:
                current_section = "fyraPositioning"
            elif 'recommendation' in lower_text:
                current_section = "recommendations"
            elif 'implementation' in lower_text or 'roadmap' in lower_text:
                current_section = "implementation"

        # Identify subsections
        elif para.style.name.startswith('Heading 2') and current_section:
            if 'sweden' in lower_text or 'swedish' in lower_text:
                current_subsection = "sweden"
            elif 'hotel' in lower_text:
                current_subsection = "hotel"
            elif 'circular' in lower_text:
                current_subsection = "circular"
            elif 'opportunit' in lower_text:
                current_subsection = "opportunities"
            elif 'challenge' in lower_text:
                current_subsection = "challenges"
            elif 'value' in lower_text:
                current_subsection = "value"
            elif 'different' in lower_text or 'unique' in lower_text:
                current_subsection = "differentiators"
            elif 'target' in lower_text:
                current_subsection = "target"
            elif 'competitive' in lower_text or 'advantage' in lower_text:
                current_subsection = "competitive"
            elif 'immediate' in lower_text or 'quick' in lower_text:
                current_subsection = "immediate"
            elif 'short' in lower_text:
                current_subsection = "short_term"
            elif 'long' in lower_text:
                current_subsection = "long_term"
            elif 'roadmap' in lower_text or 'timeline' in lower_text:
                current_subsection = "roadmap"
            elif 'milestone' in lower_text:
                current_subsection = "milestones"
            elif 'resource' in lower_text:
                current_subsection = "resources"
            elif 'metric' in lower_text or 'success' in lower_text or 'kpi' in lower_text:
                current_subsection = "metrics"

        # Add content
        elif current_section and text:
            if text.startswith('•') or text.startswith('-') or text.startswith('*'):
                item = text[1:].strip()

                if current_section == "marketAnalysis":
                    if current_subsection == "opportunities":
                        strategy["marketAnalysis"]["opportunities"].append(item)
                    elif current_subsection == "challenges":
                        strategy["marketAnalysis"]["challenges"].append(item)

                elif current_section == "fyraPositioning":
                    if current_subsection == "differentiators":
                        strategy["fyraPositioning"]["differentiators"].append(item)
                    elif current_subsection == "competitive":
                        strategy["fyraPositioning"]["competitiveAdvantages"].append(item)

                elif current_section == "recommendations":
                    if current_subsection == "immediate":
                        strategy["recommendations"]["immediate"].append(item)
                    elif current_subsection == "short_term":
                        strategy["recommendations"]["shortTerm"].append(item)
                    elif current_subsection == "long_term":
                        strategy["recommendations"]["longTerm"].append(item)

                elif current_section == "implementation":
                    if current_subsection == "roadmap":
                        strategy["implementation"]["roadmap"].append(item)
                    elif current_subsection == "milestones":
                        strategy["implementation"]["milestones"].append(item)
                    elif current_subsection == "resources":
                        strategy["implementation"]["resources"].append(item)
                    elif current_subsection == "metrics":
                        strategy["implementation"]["successMetrics"].append(item)

            else:
                # Add to overview text fields
                if current_section == "marketAnalysis":
                    if current_subsection == "sweden":
                        strategy["marketAnalysis"]["swedenOverview"] += " " + text
                    elif current_subsection == "hotel":
                        strategy["marketAnalysis"]["hotelSector"] += " " + text
                    elif current_subsection == "circular":
                        strategy["marketAnalysis"]["circularEconomyTrends"] += " " + text

                elif current_section == "fyraPositioning":
                    if current_subsection == "value":
                        strategy["fyraPositioning"]["valueProposition"] += " " + text
                    elif current_subsection == "target":
                        strategy["fyraPositioning"]["targetMarket"] += " " + text

    # Clean up text fields
    strategy["marketAnalysis"]["swedenOverview"] = clean_text(strategy["marketAnalysis"]["swedenOverview"])
    strategy["marketAnalysis"]["hotelSector"] = clean_text(strategy["marketAnalysis"]["hotelSector"])
    strategy["marketAnalysis"]["circularEconomyTrends"] = clean_text(strategy["marketAnalysis"]["circularEconomyTrends"])
    strategy["fyraPositioning"]["valueProposition"] = clean_text(strategy["fyraPositioning"]["valueProposition"])
    strategy["fyraPositioning"]["targetMarket"] = clean_text(strategy["fyraPositioning"]["targetMarket"])

    print(f"  - Extracted {len(strategy['marketAnalysis']['opportunities'])} opportunities")
    print(f"  - Extracted {len(strategy['recommendations']['immediate'])} immediate recommendations")
    print(f"  - Extracted {len(strategy['implementation']['roadmap'])} roadmap items")

    return strategy

def main():
    print("=" * 80)
    print("COMPREHENSIVE FYRA CIRCULAR PLATFORM DATA EXTRACTION")
    print("=" * 80)

    # Extract all data
    suppliers = extract_suppliers(DOCS["suppliers"])
    consultants = extract_consultants(DOCS["consultants"])
    case_studies = extract_case_studies(DOCS["hotels"])
    regulatory = extract_regulatory(DOCS["regulatory"])
    strategy = extract_strategy(DOCS["strategy"])

    # Save to JSON files
    print("\n" + "=" * 80)
    print("SAVING JSON FILES")
    print("=" * 80)

    output_files = {
        "suppliers.json": suppliers,
        "consultants.json": consultants,
        "case_studies.json": case_studies,
        "regulatory.json": regulatory,
        "strategy.json": strategy
    }

    for filename, data in output_files.items():
        output_path = DATA_DIR / filename
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"✓ Saved: {filename}")

        if isinstance(data, list):
            print(f"  - {len(data)} entries")
        elif isinstance(data, dict):
            total_items = sum(len(v) if isinstance(v, list) else 1 for v in data.values())
            print(f"  - {len(data)} sections, ~{total_items} items")

    # Generate summary
    print("\n" + "=" * 80)
    print("EXTRACTION SUMMARY")
    print("=" * 80)
    print(f"\nSuppliers: {len(suppliers)} companies extracted")
    for s in suppliers[:3]:
        print(f"  - {s['name']}: {len(s['services'])} services, {len(s['certifications'])} certifications")

    print(f"\nConsultants: {len(consultants)} firms extracted")
    for c in consultants[:3]:
        print(f"  - {c['name']}: {len(c['services'])} services, {len(c['specializations'])} specializations")

    print(f"\nCase Studies: {len(case_studies)} hotels extracted")
    for cs in case_studies[:3]:
        print(f"  - {cs['hotelName']}: {cs['year']}, {len(cs['circularElements'])} circular elements")

    print(f"\nRegulatory: {len(regulatory)} main sections")
    print(f"  - Fire Safety: {len(regulatory['fireSafety']['bbrRequirements'])} BBR requirements")
    print(f"  - Building Codes: {len(regulatory['buildingCodes']['hotelSpecific'])} hotel-specific items")
    print(f"  - Implementation: {len(regulatory['implementation']['tips'])} practical tips")

    print(f"\nStrategy: {len(strategy)} main sections")
    print(f"  - Market Opportunities: {len(strategy['marketAnalysis']['opportunities'])} identified")
    print(f"  - Fyra Differentiators: {len(strategy['fyraPositioning']['differentiators'])} listed")
    print(f"  - Implementation Roadmap: {len(strategy['implementation']['roadmap'])} steps")

    print("\n" + "=" * 80)
    print("EXTRACTION COMPLETE!")
    print("=" * 80)
    print(f"\nAll files saved to: {DATA_DIR}")

if __name__ == "__main__":
    main()
