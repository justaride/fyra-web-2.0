#!/usr/bin/env python3
"""
Extract clean, structured data from NCC Word documents for Fyra website
"""
import json
import re
from docx import Document
from pathlib import Path

def extract_suppliers():
    """Extract B2B reuse operators from document 1"""
    doc = Document("1. COMPREHENSIVE B2B REUSE OPERATORS EXTRACTION.docx")
    suppliers = []

    current_supplier = None

    for para in doc.paragraphs:
        text = para.text.strip()

        # Skip empty lines
        if not text:
            continue

        # Check for company headers (usually all caps or bold)
        if text.isupper() and len(text) < 100:
            if current_supplier:
                suppliers.append(current_supplier)
            current_supplier = {
                "name": text.replace("⭐", "").strip(),
                "location": "",
                "description": "",
                "services": [],
                "capabilities": [],
                "contact": {},
                "website": "",
                "hospitalityReady": False,
                "volumeCapacity": "",
                "logistics": "",
                "certifications": []
            }
        elif current_supplier:
            # Extract structured information
            if "website:" in text.lower() or "http" in text.lower():
                urls = re.findall(r'https?://[^\s]+', text)
                if urls:
                    current_supplier["website"] = urls[0].strip()
            elif "phone:" in text.lower() or "+46" in text:
                phone = re.search(r'\+46[^a-zA-Z]+', text)
                if phone:
                    current_supplier["contact"]["phone"] = phone.group().strip()
            elif "email:" in text.lower() or "@" in text:
                email = re.search(r'[\w\.-]+@[\w\.-]+', text)
                if email:
                    current_supplier["contact"]["email"] = email.group().strip()
            elif "location:" in text.lower() or "stockholm" in text.lower() or "sweden" in text.lower():
                current_supplier["location"] = text.replace("Location:", "").strip()
            else:
                # Add to description
                if len(text) > 20 and current_supplier["description"]:
                    current_supplier["description"] += " " + text
                elif len(text) > 20:
                    current_supplier["description"] = text

    # Add last supplier
    if current_supplier:
        suppliers.append(current_supplier)

    # Filter out non-supplier entries (actions, etc.)
    clean_suppliers = [s for s in suppliers if not any(word in s["name"].lower()
                       for word in ["call", "develop", "create", "request", "priority"])]

    return clean_suppliers

def extract_consultants():
    """Extract PM firms from document 3"""
    doc = Document("3. SWEDISH PROJECT MANAGERS & TECHNICAL CONSULTANTS FOR WHOLE-PROJECT HOSPITALITY RENOVATION.docx")
    consultants = []

    current_consultant = None

    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        # Look for company names (often in headers or bold)
        if text.isupper() and len(text) < 100 and "AB" in text or "SWECO" in text.upper():
            if current_consultant:
                consultants.append(current_consultant)
            current_consultant = {
                "name": text,
                "services": [],
                "hospitalityExperience": False,
                "circularEconomyExperience": False,
                "contact": {},
                "website": "",
                "description": ""
            }
        elif current_consultant:
            # Extract info
            if "website:" in text.lower() or "http" in text.lower():
                urls = re.findall(r'https?://[^\s]+', text)
                if urls:
                    current_consultant["website"] = urls[0].strip()
            elif len(text) > 20:
                current_consultant["description"] = current_consultant.get("description", "") + " " + text

    if current_consultant:
        consultants.append(current_consultant)

    return consultants

def extract_case_studies():
    """Extract hotel case studies from document 4"""
    doc = Document("4. FRONTRUNNER HOTELS - COMPLETE EXTRACTION.docx")
    cases = []

    current_case = None

    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        # Look for hotel names
        if "hotel" in text.lower() or "renovation" in text.lower():
            if len(text) < 150 and (text.isupper() or para.style.name.startswith('Heading')):
                if current_case:
                    cases.append(current_case)
                current_case = {
                    "name": text,
                    "location": "",
                    "year": "",
                    "circularElements": [],
                    "impact": {},
                    "description": ""
                }
        elif current_case:
            if len(text) > 20:
                current_case["description"] = current_case.get("description", "") + " " + text

                # Extract location
                if any(city in text for city in ["Stockholm", "Gothenburg", "Malmö", "Copenhagen", "Oslo"]):
                    current_case["location"] = text

    if current_case:
        cases.append(current_case)

    return cases[:30]  # Limit to top 30 most detailed cases

def main():
    print("Extracting clean data from NCC documents...")

    # Extract all data
    suppliers = extract_suppliers()
    consultants = extract_consultants()
    cases = extract_case_studies()

    # Save to JSON
    with open("fyra-circular-platform/data/suppliers_clean.json", "w", encoding="utf-8") as f:
        json.dump(suppliers, f, indent=2, ensure_ascii=False)

    with open("fyra-circular-platform/data/consultants_clean.json", "w", encoding="utf-8") as f:
        json.dump(consultants, f, indent=2, ensure_ascii=False)

    with open("fyra-circular-platform/data/cases_clean.json", "w", encoding="utf-8") as f:
        json.dump(cases, f, indent=2, ensure_ascii=False)

    print(f"✓ Extracted {len(suppliers)} suppliers")
    print(f"✓ Extracted {len(consultants)} consultants")
    print(f"✓ Extracted {len(cases)} case studies")
    print("\nFiles created in fyra-circular-platform/data/")

if __name__ == "__main__":
    main()
