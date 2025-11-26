"""
Production-ready extraction with comprehensive table parsing
"""

import json
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import re
from pathlib import Path
from collections import defaultdict

def clean_text(text):
    """Clean and normalize text"""
    if not text:
        return ""
    return ' '.join(text.strip().split())

def extract_all_tables(doc):
    """Extract all tables from document"""
    tables_data = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(clean_text(cell.text))
            table_data.append(row_data)
        if table_data:
            tables_data.append(table_data)
    return tables_data

def parse_supplier_table(table_data):
    """Parse supplier/operator tables"""
    if not table_data or len(table_data) < 2:
        return []

    suppliers = []
    headers = [h.lower() for h in table_data[0]]

    for row in table_data[1:]:
        if len(row) < 2 or not row[0]:
            continue

        supplier = {
            'type': 'B2B Reuse Operator',
            'services': [],
            'capabilities': [],
            'certifications': []
        }

        for i, value in enumerate(row):
            if i >= len(headers):
                break
            header = headers[i]
            value = clean_text(value)

            if not value or value == '-':
                continue

            # Map headers to fields
            if 'name' in header or 'company' in header or 'operator' in header or 'företag' in header:
                supplier['name'] = value
            elif 'location' in header or 'stad' in header or 'city' in header:
                supplier['location'] = value
            elif 'focus' in header or 'specialty' in header or 'special' in header:
                supplier['specialty'] = value
            elif 'service' in header or 'tjänst' in header:
                supplier['services'].append(value)
            elif 'capability' in header or 'capabilit' in header:
                supplier['capabilities'].append(value)
            elif 'volume' in header or 'capacity' in header or 'volym' in header:
                supplier['volumeCapacity'] = value
            elif 'quality' in header or 'kvalitet' in header:
                supplier['qualityAssurance'] = value
            elif 'cert' in header:
                supplier['certifications'].append(value)
            elif 'contact' in header or 'kontakt' in header:
                supplier['contact'] = value
            elif 'email' in header or 'e-post' in header:
                supplier['email'] = value
            elif 'website' in header or 'webb' in header:
                supplier['website'] = value
            elif 'phone' in header or 'tel' in header:
                supplier['phone'] = value
            elif 'hospitality' in header:
                supplier['hospitalityReadiness'] = value
            elif 'lead time' in header or 'ledtid' in header:
                supplier['leadTime'] = value
            elif 'pricing' in header or 'pris' in header:
                supplier['pricing'] = value
            else:
                # Store unmapped fields
                field_name = header.replace(' ', '_').replace('-', '_')
                if field_name:
                    supplier[field_name] = value

        if supplier.get('name') and len(supplier['name']) > 2:
            suppliers.append(supplier)

    return suppliers

def extract_suppliers_comprehensive(doc_paths):
    """Extract suppliers from multiple documents comprehensively"""
    all_suppliers = {}

    for doc_path in doc_paths:
        doc = Document(doc_path)

        # Extract from tables
        tables = extract_all_tables(doc)
        for table in tables:
            suppliers = parse_supplier_table(table)
            for supplier in suppliers:
                name = supplier.get('name')
                if name:
                    if name in all_suppliers:
                        # Merge data
                        for key, value in supplier.items():
                            if key not in all_suppliers[name] or not all_suppliers[name][key]:
                                all_suppliers[name][key] = value
                            elif isinstance(value, list) and isinstance(all_suppliers[name][key], list):
                                all_suppliers[name][key].extend(value)
                    else:
                        all_suppliers[name] = supplier

        # Also extract from structured paragraphs
        current_supplier = None
        in_supplier_section = False

        for para in doc.paragraphs:
            text = clean_text(para.text)
            if not text:
                continue

            # Detect supplier sections
            if re.match(r'^\d+\.\s+[A-ZÅÄÖ]', text) and len(text) < 100:
                if current_supplier and current_supplier.get('name'):
                    name = current_supplier['name']
                    if name in all_suppliers:
                        for key, value in current_supplier.items():
                            if key not in all_suppliers[name] or not all_suppliers[name][key]:
                                all_suppliers[name][key] = value
                    else:
                        all_suppliers[name] = current_supplier

                # Start new supplier
                name = re.sub(r'^\d+\.\s+', '', text)
                name = re.sub(r'\s*\(.*?\)', '', name)
                current_supplier = {
                    'name': name.strip(),
                    'type': 'B2B Reuse Operator',
                    'services': [],
                    'capabilities': [],
                    'certifications': []
                }
                in_supplier_section = True

            elif in_supplier_section and current_supplier and ':' in text:
                parts = text.split(':', 1)
                key = parts[0].strip().lower()
                value = parts[1].strip() if len(parts) > 1 else ""

                if value and len(value) > 1:
                    if 'location' in key or 'stad' in key:
                        current_supplier['location'] = value
                    elif 'website' in key:
                        current_supplier['website'] = value
                    elif 'email' in key:
                        current_supplier['email'] = value
                    elif 'phone' in key or 'tel' in key:
                        current_supplier['phone'] = value
                    elif 'contact' in key:
                        current_supplier['contact'] = value
                    elif 'focus' in key or 'specialty' in key:
                        current_supplier['specialty'] = value
                    elif 'volume' in key:
                        current_supplier['volumeCapacity'] = value
                    elif 'quality' in key:
                        current_supplier['qualityAssurance'] = value
                    elif 'hospitality readiness' in key:
                        current_supplier['hospitalityReadiness'] = value
                    elif 'best for' in key:
                        current_supplier['bestFor'] = value
                    elif 'key strength' in key:
                        current_supplier['keyStrength'] = value
                    elif 'key gap' in key:
                        current_supplier['keyGap'] = value

        # Add last supplier
        if current_supplier and current_supplier.get('name'):
            name = current_supplier['name']
            if name in all_suppliers:
                for key, value in current_supplier.items():
                    if key not in all_suppliers[name] or not all_suppliers[name][key]:
                        all_suppliers[name][key] = value
            else:
                all_suppliers[name] = current_supplier

    # Filter and clean
    result = []
    for name, supplier in all_suppliers.items():
        if len(name) > 2 and not any(skip in name.upper() for skip in
            ['TIER', 'COMPREHENSIVE', 'CONTRACT-LEVEL', 'NOTE:', 'CRITICAL']):
            result.append(supplier)

    return result

def extract_consultants_comprehensive(doc_path):
    """Extract consultants comprehensively"""
    doc = Document(doc_path)
    consultants = {}

    # Extract from tables
    tables = extract_all_tables(doc)
    for table in tables:
        if len(table) < 2:
            continue

        headers = [h.lower() for h in table[0]]

        for row in table[1:]:
            if len(row) < 2 or not row[0]:
                continue

            consultant = {
                'type': 'Project Manager / Technical Consultant',
                'projects': [],
                'certifications': [],
                'expertise': []
            }

            for i, value in enumerate(row):
                if i >= len(headers):
                    break
                header = headers[i]
                value = clean_text(value)

                if not value or value == '-':
                    continue

                if 'name' in header or 'company' in header or 'företag' in header:
                    consultant['name'] = value
                elif 'location' in header or 'stad' in header:
                    consultant['location'] = value
                elif 'specialty' in header or 'special' in header:
                    consultant['specialty'] = value
                elif 'expertise' in header:
                    consultant['expertise'].append(value)
                elif 'experience' in header or 'track record' in header:
                    consultant['experience'] = value
                elif 'cert' in header:
                    consultant['certifications'].append(value)
                elif 'contact' in header:
                    consultant['contact'] = value
                elif 'email' in header:
                    consultant['email'] = value
                elif 'website' in header:
                    consultant['website'] = value
                elif 'phone' in header:
                    consultant['phone'] = value
                elif 'hospitality' in header:
                    consultant['hospitalityExperience'] = value
                elif 'circular' in header:
                    consultant['circularEconomyExperience'] = value
                else:
                    field_name = header.replace(' ', '_').replace('-', '_')
                    if field_name:
                        consultant[field_name] = value

            if consultant.get('name'):
                consultants[consultant['name']] = consultant

    # Extract from paragraphs
    current_consultant = None

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        if re.match(r'^\d+\.\s+[A-ZÅÄÖ]', text) and len(text) < 150:
            if current_consultant and current_consultant.get('name'):
                name = current_consultant['name']
                if name in consultants:
                    for key, value in current_consultant.items():
                        if key not in consultants[name] or not consultants[name][key]:
                            consultants[name][key] = value
                else:
                    consultants[name] = current_consultant

            name = re.sub(r'^\d+\.\s+', '', text)
            name = re.sub(r'\s*-\s*PRIORITY.*', '', name)
            name = re.sub(r'\s*-\s*RANKING.*', '', name)
            current_consultant = {
                'name': name.strip(),
                'type': 'Project Manager / Technical Consultant',
                'projects': [],
                'certifications': [],
                'expertise': []
            }

        elif current_consultant and ':' in text:
            parts = text.split(':', 1)
            key = parts[0].strip().lower()
            value = parts[1].strip() if len(parts) > 1 else ""

            if value and len(value) > 1:
                if 'location' in key:
                    current_consultant['location'] = value
                elif 'website' in key:
                    current_consultant['website'] = value
                elif 'contact' in key or 'email' in key:
                    current_consultant['contact'] = value
                elif 'phone' in key:
                    current_consultant['phone'] = value
                elif 'hospitality' in key:
                    current_consultant['hospitalityExperience'] = value
                elif 'circular' in key:
                    current_consultant['circularEconomyExperience'] = value
                elif 'key strength' in key:
                    current_consultant['keyStrength'] = value

    if current_consultant and current_consultant.get('name'):
        name = current_consultant['name']
        if name in consultants:
            for key, value in current_consultant.items():
                if key not in consultants[name] or not consultants[name][key]:
                    consultants[name][key] = value
        else:
            consultants[name] = current_consultant

    # Filter
    result = []
    for name, consultant in consultants.items():
        if len(name) > 2 and not any(skip in name.upper() for skip in
            ['TIER', 'EXECUTIVE', 'CONTACT &', 'SWEDISH PROJECT']):
            result.append(consultant)

    return result

def extract_case_studies_comprehensive(doc_path):
    """Extract case studies comprehensively"""
    doc = Document(doc_path)
    case_studies = {}

    # Extract from tables
    tables = extract_all_tables(doc)
    for table in tables:
        if len(table) < 2:
            continue

        headers = [h.lower() for h in table[0]]

        for row in table[1:]:
            if len(row) < 2 or not row[0]:
                continue

            case = {
                'type': 'Hotel Case Study',
                'circularElements': [],
                'materials': []
            }

            for i, value in enumerate(row):
                if i >= len(headers):
                    break
                header = headers[i]
                value = clean_text(value)

                if not value or value == '-':
                    continue

                if 'hotel' in header or 'name' in header:
                    case['name'] = value
                elif 'location' in header or 'stad' in header or 'address' in header:
                    case['location'] = value
                elif 'size' in header or 'rooms' in header:
                    case['size'] = value
                elif 'year' in header or 'opening' in header:
                    case['year'] = value
                elif 'chain' in header or 'brand' in header:
                    case['chain'] = value
                elif 'category' in header:
                    case['category'] = value
                elif 'scope' in header:
                    case['scope'] = value
                elif 'circular' in header or 'reuse' in header:
                    case['circularPercentage'] = value
                elif 'co2' in header or 'carbon' in header:
                    case['co2Savings'] = value
                elif 'financial' in header or 'cost' in header:
                    case['financialImpact'] = value
                elif 'architect' in header:
                    case['architect'] = value
                elif 'contractor' in header:
                    case['contractor'] = value
                elif 'website' in header:
                    case['website'] = value
                elif 'phone' in header:
                    case['phone'] = value
                else:
                    field_name = header.replace(' ', '_').replace('-', '_')
                    if field_name:
                        case[field_name] = value

            if case.get('name'):
                case_studies[case['name']] = case

    # Extract from paragraphs
    current_case = None

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        # Detect hotel names (often in all caps or specific patterns)
        if (para.style.name.startswith('Heading') or
            (para.runs and para.runs[0].bold and len(text) < 100 and
             ('HOTEL' in text.upper() or 'BY NOBIS' in text.upper() or
              re.match(r'^[A-ZÅÄÖ\s&-]+$', text)))):

            if current_case and current_case.get('name'):
                name = current_case['name']
                if name in case_studies:
                    for key, value in current_case.items():
                        if key not in case_studies[name] or not case_studies[name][key]:
                            case_studies[name][key] = value
                else:
                    case_studies[name] = current_case

            name = re.sub(r'\s*-\s*RELEVANCE SCORE:.*', '', text)
            current_case = {
                'name': name.strip(),
                'type': 'Hotel Case Study',
                'circularElements': [],
                'materials': []
            }

        elif current_case and ':' in text:
            parts = text.split(':', 1)
            key = parts[0].strip().lower()
            value = parts[1].strip() if len(parts) > 1 else ""

            if value and len(value) > 1:
                if 'name:' in text.lower() and 'hotel' in value.lower():
                    current_case['name'] = value
                elif 'location' in key or 'address' in key:
                    current_case['location'] = value
                elif 'size' in key or 'rooms' in key:
                    current_case['size'] = value
                elif 'year' in key or 'opening' in key:
                    current_case['year'] = value
                elif 'chain' in key:
                    current_case['chain'] = value
                elif 'category' in key:
                    current_case['category'] = value
                elif 'scope' in key:
                    current_case['scope'] = value
                elif 'co2' in key or 'carbon' in key:
                    current_case['co2Savings'] = value
                elif 'financial' in key or 'cost' in key:
                    current_case['financialImpact'] = value
                elif 'waste' in key:
                    current_case['wasteReduction'] = value
                elif 'architect' in key:
                    current_case['architect'] = value
                elif 'contractor' in key:
                    current_case['contractor'] = value
                elif 'website' in key:
                    current_case['website'] = value
                elif 'phone' in key:
                    current_case['phone'] = value

    if current_case and current_case.get('name'):
        name = current_case['name']
        if name in case_studies:
            for key, value in current_case.items():
                if key not in case_studies[name] or not case_studies[name][key]:
                    case_studies[name][key] = value
        else:
            case_studies[name] = current_case

    # Filter
    result = []
    for name, case in case_studies.items():
        if len(name) > 2 and not any(skip in name.upper() for skip in
            ['PROFILE', 'CIRCULAR ECONOMY', 'QUANTIFIED', 'FRONTRUNNER']):
            result.append(case)

    return result

def extract_regulations_comprehensive(doc_path):
    """Extract regulations comprehensively"""
    doc = Document(doc_path)

    regulations = {
        'fireSafety': {'title': 'Fire Safety Requirements', 'items': []},
        'buildingCodes': {'title': 'Building Codes (BBR)', 'items': []},
        'bvb': {'title': 'BVB Standards', 'items': []},
        'materialStandards': {'title': 'Material Standards & CE Marking', 'items': []},
        'documentation': {'title': 'Documentation Requirements', 'items': []},
        'compliance': {'title': 'Compliance & Certification', 'items': []}
    }

    current_category = 'documentation'
    current_item = None

    # Extract from tables
    tables = extract_all_tables(doc)
    for table in tables:
        if len(table) < 2:
            continue

        for row in table:
            if len(row) < 1:
                continue

            # Determine category from content
            row_text = ' '.join(row).lower()
            if 'fire' in row_text or 'brand' in row_text:
                category = 'fireSafety'
            elif 'bbr' in row_text or 'boverket' in row_text:
                category = 'buildingCodes'
            elif 'bvb' in row_text:
                category = 'bvb'
            elif 'ce marking' in row_text or 'material standard' in row_text:
                category = 'materialStandards'
            elif 'document' in row_text:
                category = 'documentation'
            elif 'compli' in row_text or 'certif' in row_text:
                category = 'compliance'
            else:
                category = current_category

            if len(row) >= 2:
                regulations[category]['items'].append({
                    'title': row[0],
                    'description': ' | '.join(row[1:])
                })
            else:
                regulations[category]['items'].append({
                    'title': row[0],
                    'description': ''
                })

    # Extract from paragraphs
    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text or len(text) < 5:
            continue

        # Detect category headers
        text_lower = text.lower()
        if 'fire safety' in text_lower or 'brandsäkerhet' in text_lower:
            current_category = 'fireSafety'
            current_item = None
        elif 'bbr' in text_lower or 'building code' in text_lower:
            current_category = 'buildingCodes'
            current_item = None
        elif 'bvb' in text_lower and 'standard' in text_lower:
            current_category = 'bvb'
            current_item = None
        elif 'material standard' in text_lower or 'ce marking' in text_lower:
            current_category = 'materialStandards'
            current_item = None
        elif 'documentation' in text_lower:
            current_category = 'documentation'
            current_item = None
        elif 'compliance' in text_lower or 'certification' in text_lower:
            current_category = 'compliance'
            current_item = None

        # Store items
        if para.style.name.startswith('Heading') or (para.runs and para.runs[0].bold):
            if current_item and current_item.get('title'):
                regulations[current_category]['items'].append(current_item)
            current_item = {'title': text, 'description': ''}
        elif current_item:
            current_item['description'] += ' ' + text
        elif len(text) > 10:
            regulations[current_category]['items'].append({
                'title': text[:100] + ('...' if len(text) > 100 else ''),
                'description': text if len(text) > 100 else ''
            })

    if current_item and current_item.get('title'):
        regulations[current_category]['items'].append(current_item)

    # Clean descriptions
    for category in regulations:
        for item in regulations[category]['items']:
            item['description'] = item['description'].strip()

    return regulations

def extract_implementation_comprehensive(doc_path):
    """Extract implementation roadmap comprehensively"""
    doc = Document(doc_path)

    implementation = {
        'executiveSummary': '',
        'marketMaturity': {},
        'strengths': [],
        'weaknesses': [],
        'roadmap': [],
        'marketAnalysis': {},
        'recommendations': [],
        'kpis': [],
        'timeline': []
    }

    current_section = 'executiveSummary'
    current_text = []

    # Extract from tables
    tables = extract_all_tables(doc)
    for table in tables:
        if len(table) < 2:
            continue

        headers = [h.lower() for h in table[0]]
        table_data = []

        for row in table[1:]:
            row_data = {}
            for i, value in enumerate(row):
                if i < len(headers) and headers[i]:
                    row_data[headers[i].replace(' ', '_')] = value
            if row_data:
                table_data.append(row_data)

        # Categorize table data
        header_text = ' '.join(headers).lower()
        if 'phase' in header_text or 'roadmap' in header_text:
            implementation['roadmap'].extend(table_data)
        elif 'timeline' in header_text or 'schedule' in header_text:
            implementation['timeline'].extend(table_data)
        elif 'kpi' in header_text or 'metric' in header_text:
            implementation['kpis'].extend(table_data)
        else:
            implementation['recommendations'].extend(table_data)

    # Extract from paragraphs
    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        text_lower = text.lower()

        # Detect sections
        if 'executive summary' in text_lower:
            if current_text:
                implementation[current_section] = ' '.join(current_text)
            current_section = 'executiveSummary'
            current_text = []
        elif 'market maturity' in text_lower or 'readiness' in text_lower:
            if current_text:
                if current_section in ['strengths', 'weaknesses', 'recommendations']:
                    implementation[current_section].extend(current_text)
                else:
                    implementation[current_section] = ' '.join(current_text)
            current_section = 'marketMaturity'
            current_text = []
        elif 'strength' in text_lower and para.style.name.startswith('Heading'):
            if current_text:
                implementation[current_section] = ' '.join(current_text)
            current_section = 'strengths'
            current_text = []
        elif 'weakness' in text_lower or 'gap' in text_lower:
            if current_text:
                if isinstance(implementation[current_section], list):
                    implementation[current_section].extend(current_text)
                else:
                    implementation[current_section] = ' '.join(current_text)
            current_section = 'weaknesses'
            current_text = []
        elif 'recommendation' in text_lower and para.style.name.startswith('Heading'):
            if current_text:
                if isinstance(implementation[current_section], list):
                    implementation[current_section].extend(current_text)
                else:
                    implementation[current_section] = ' '.join(current_text)
            current_section = 'recommendations'
            current_text = []
        else:
            current_text.append(text)

    # Add remaining text
    if current_text:
        if isinstance(implementation[current_section], list):
            implementation[current_section].extend(current_text)
        else:
            implementation[current_section] = ' '.join(current_text)

    return implementation

def main():
    """Main extraction process"""
    base_path = Path('/Users/gabrielboen/Downloads/drive-download-20251108T110451Z-1-001')

    print("="*70)
    print("COMPREHENSIVE DATA EXTRACTION")
    print("="*70)

    # Extract suppliers
    print("\n[1/5] Extracting B2B Reuse Operators...")
    suppliers = extract_suppliers_comprehensive([
        base_path / '1. COMPREHENSIVE B2B REUSE OPERATORS EXTRACTION.docx',
        base_path / '2. FYRA B2B OPERATOR CAPABILITY MATRIX.docx'
    ])

    with open(base_path / 'suppliers.json', 'w', encoding='utf-8') as f:
        json.dump(suppliers, f, indent=2, ensure_ascii=False)
    print(f"      ✓ {len(suppliers)} suppliers extracted → suppliers.json")

    # Extract consultants
    print("\n[2/5] Extracting Project Managers & Consultants...")
    consultants = extract_consultants_comprehensive(
        base_path / '3. SWEDISH PROJECT MANAGERS & TECHNICAL CONSULTANTS FOR WHOLE-PROJECT HOSPITALITY RENOVATION.docx'
    )

    with open(base_path / 'consultants.json', 'w', encoding='utf-8') as f:
        json.dump(consultants, f, indent=2, ensure_ascii=False)
    print(f"      ✓ {len(consultants)} consultants extracted → consultants.json")

    # Extract case studies
    print("\n[3/5] Extracting Hotel Case Studies...")
    case_studies = extract_case_studies_comprehensive(
        base_path / '4. FRONTRUNNER HOTELS - COMPLETE EXTRACTION.docx'
    )

    with open(base_path / 'caseStudies.json', 'w', encoding='utf-8') as f:
        json.dump(case_studies, f, indent=2, ensure_ascii=False)
    print(f"      ✓ {len(case_studies)} case studies extracted → caseStudies.json")

    # Extract regulations
    print("\n[4/5] Extracting Swedish Regulations...")
    regulations = extract_regulations_comprehensive(
        base_path / '5. PRACTICAL REGULATORY GUIDE_ REUSED MATERIALS IN SWEDISH HOTEL RENOVATIONS.docx'
    )

    with open(base_path / 'regulations.json', 'w', encoding='utf-8') as f:
        json.dump(regulations, f, indent=2, ensure_ascii=False)

    total_regs = sum(len(v['items']) for v in regulations.values())
    print(f"      ✓ {total_regs} regulations across {len(regulations)} categories → regulations.json")

    # Extract implementation
    print("\n[5/5] Extracting Strategic Roadmap...")
    implementation = extract_implementation_comprehensive(
        base_path / '6. STRATEGIC IMPLEMENTATION ROADMAP.docx'
    )

    with open(base_path / 'implementation.json', 'w', encoding='utf-8') as f:
        json.dump(implementation, f, indent=2, ensure_ascii=False)
    print(f"      ✓ Implementation roadmap extracted → implementation.json")

    # Create summary
    summary = {
        'extractionDate': '2025-11-08',
        'version': '3.0-production',
        'sourceDocuments': [
            '1. COMPREHENSIVE B2B REUSE OPERATORS EXTRACTION.docx',
            '2. FYRA B2B OPERATOR CAPABILITY MATRIX.docx',
            '3. SWEDISH PROJECT MANAGERS & TECHNICAL CONSULTANTS.docx',
            '4. FRONTRUNNER HOTELS - COMPLETE EXTRACTION.docx',
            '5. PRACTICAL REGULATORY GUIDE.docx',
            '6. STRATEGIC IMPLEMENTATION ROADMAP.docx'
        ],
        'outputFiles': {
            'suppliers.json': {
                'count': len(suppliers),
                'description': 'B2B reuse operators with contact info, capabilities, and service details'
            },
            'consultants.json': {
                'count': len(consultants),
                'description': 'Project managers and technical consultants for hotel renovations'
            },
            'caseStudies.json': {
                'count': len(case_studies),
                'description': 'Real hotel case studies with circular economy implementation details'
            },
            'regulations.json': {
                'count': total_regs,
                'categories': list(regulations.keys()),
                'description': 'Swedish regulatory requirements organized by category'
            },
            'implementation.json': {
                'description': 'Strategic roadmap, market analysis, and recommendations'
            }
        },
        'dataStructure': {
            'suppliers': list(suppliers[0].keys()) if suppliers else [],
            'consultants': list(consultants[0].keys()) if consultants else [],
            'caseStudies': list(case_studies[0].keys()) if case_studies else [],
            'regulations': {
                'structure': 'Object with categories (fireSafety, buildingCodes, bvb, materialStandards, documentation, compliance)',
                'itemStructure': 'Each category has title and array of items with title/description'
            },
            'implementation': list(implementation.keys())
        }
    }

    with open(base_path / 'extraction_summary.json', 'w', encoding='utf-8') as f:
        json.dump(summary, f, indent=2, ensure_ascii=False)

    print("\n" + "="*70)
    print("EXTRACTION COMPLETE!")
    print("="*70)
    print(f"\nTotal Data Extracted:")
    print(f"  • {len(suppliers)} B2B Reuse Operators")
    print(f"  • {len(consultants)} Project Managers/Technical Consultants")
    print(f"  • {len(case_studies)} Hotel Case Studies")
    print(f"  • {total_regs} Regulatory Requirements")
    print(f"  • Strategic Implementation Roadmap")
    print(f"\nJSON Files Created:")
    print(f"  1. suppliers.json")
    print(f"  2. consultants.json")
    print(f"  3. caseStudies.json")
    print(f"  4. regulations.json")
    print(f"  5. implementation.json")
    print(f"  6. extraction_summary.json")
    print(f"\nLocation: {base_path}")
    print("\n✓ All data is clean, structured, and ready for Next.js integration!")

if __name__ == '__main__':
    main()
