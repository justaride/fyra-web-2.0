"""
Extract data from Word documents and create structured JSON files for Next.js website
"""

import json
from docx import Document
import re
from pathlib import Path

def clean_text(text):
    """Clean and normalize text"""
    if not text:
        return ""
    return text.strip().replace('\n', ' ').replace('\r', '')

def extract_suppliers(doc_path):
    """Extract B2B reuse operators from document 1 & 2"""
    doc = Document(doc_path)
    suppliers = []
    current_supplier = {}

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        # Detect new supplier entry (typically starts with company name in bold or numbered)
        if para.runs and para.runs[0].bold and len(text) < 100:
            if current_supplier.get('name'):
                suppliers.append(current_supplier)
            current_supplier = {'name': text}
        else:
            # Parse key-value patterns
            if ':' in text:
                parts = text.split(':', 1)
                key = parts[0].strip().lower()
                value = parts[1].strip() if len(parts) > 1 else ""

                # Map common fields
                field_mapping = {
                    'location': 'location',
                    'lokalisering': 'location',
                    'stad': 'location',
                    'city': 'location',
                    'services': 'services',
                    'tjänster': 'services',
                    'capabilities': 'capabilities',
                    'förmågor': 'capabilities',
                    'certifications': 'certifications',
                    'certifieringar': 'certifications',
                    'contact': 'contact',
                    'kontakt': 'contact',
                    'email': 'email',
                    'website': 'website',
                    'webbplats': 'website',
                    'phone': 'phone',
                    'telefon': 'phone',
                    'volume capacity': 'volumeCapacity',
                    'volymkapacitet': 'volumeCapacity',
                    'quality assurance': 'qualityAssurance',
                    'kvalitetssäkring': 'qualityAssurance',
                    'logistics': 'logistics',
                    'logistik': 'logistics',
                    'specialty': 'specialty',
                    'specialitet': 'specialty',
                    'materials': 'materials',
                    'material': 'materials',
                    'scope': 'scope',
                    'omfattning': 'scope',
                    'pricing': 'pricing',
                    'prissättning': 'pricing',
                    'lead time': 'leadTime',
                    'ledtid': 'leadTime',
                }

                for key_pattern, field_name in field_mapping.items():
                    if key_pattern in key:
                        current_supplier[field_name] = value
                        break

    # Add last supplier
    if current_supplier.get('name'):
        suppliers.append(current_supplier)

    # Also extract from tables
    for table in doc.tables:
        headers = [clean_text(cell.text).lower() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            supplier = {}
            for idx, cell in enumerate(row.cells):
                value = clean_text(cell.text)
                if idx < len(headers) and headers[idx]:
                    # Map header to field
                    header = headers[idx]
                    if 'name' in header or 'företag' in header or 'company' in header:
                        supplier['name'] = value
                    elif 'location' in header or 'stad' in header or 'city' in header:
                        supplier['location'] = value
                    elif 'service' in header or 'tjänst' in header:
                        supplier['services'] = value
                    elif 'capabilit' in header or 'förmåg' in header:
                        supplier['capabilities'] = value
                    elif 'cert' in header:
                        supplier['certifications'] = value
                    elif 'contact' in header or 'kontakt' in header:
                        supplier['contact'] = value
                    elif 'email' in header or 'e-post' in header:
                        supplier['email'] = value
                    elif 'website' in header or 'webb' in header:
                        supplier['website'] = value
                    elif 'volume' in header or 'volym' in header:
                        supplier['volumeCapacity'] = value
                    elif 'quality' in header or 'kvalitet' in header:
                        supplier['qualityAssurance'] = value
                    elif 'logistics' in header or 'logistik' in header:
                        supplier['logistics'] = value
                    elif 'specialty' in header or 'special' in header:
                        supplier['specialty'] = value
                    elif 'material' in header:
                        supplier['materials'] = value
                    elif 'scope' in header or 'omfattning' in header:
                        supplier['scope'] = value
                    elif 'pricing' in header or 'pris' in header:
                        supplier['pricing'] = value
                    elif 'lead' in header or 'ledtid' in header:
                        supplier['leadTime'] = value
                    else:
                        supplier[header.replace(' ', '_')] = value

            if supplier.get('name'):
                suppliers.append(supplier)

    return suppliers

def extract_consultants(doc_path):
    """Extract project managers and technical consultants from document 3"""
    doc = Document(doc_path)
    consultants = []
    current_consultant = {}

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        # Detect new consultant entry
        if para.runs and para.runs[0].bold and len(text) < 100:
            if current_consultant.get('name'):
                consultants.append(current_consultant)
            current_consultant = {'name': text}
        else:
            # Parse key-value patterns
            if ':' in text:
                parts = text.split(':', 1)
                key = parts[0].strip().lower()
                value = parts[1].strip() if len(parts) > 1 else ""

                field_mapping = {
                    'location': 'location',
                    'city': 'location',
                    'specialty': 'specialty',
                    'specialitet': 'specialty',
                    'expertise': 'expertise',
                    'experience': 'experience',
                    'erfarenhet': 'experience',
                    'certifications': 'certifications',
                    'contact': 'contact',
                    'email': 'email',
                    'phone': 'phone',
                    'website': 'website',
                    'projects': 'projects',
                    'projekt': 'projects',
                    'services': 'services',
                    'team size': 'teamSize',
                    'availability': 'availability',
                    'rate': 'rate',
                    'languages': 'languages',
                }

                for key_pattern, field_name in field_mapping.items():
                    if key_pattern in key:
                        current_consultant[field_name] = value
                        break

    if current_consultant.get('name'):
        consultants.append(current_consultant)

    # Extract from tables
    for table in doc.tables:
        headers = [clean_text(cell.text).lower() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            consultant = {}
            for idx, cell in enumerate(row.cells):
                value = clean_text(cell.text)
                if idx < len(headers) and headers[idx]:
                    header = headers[idx]
                    if 'name' in header or 'företag' in header:
                        consultant['name'] = value
                    elif 'location' in header or 'stad' in header:
                        consultant['location'] = value
                    elif 'specialty' in header or 'special' in header:
                        consultant['specialty'] = value
                    elif 'expertise' in header:
                        consultant['expertise'] = value
                    elif 'experience' in header or 'erfaren' in header:
                        consultant['experience'] = value
                    elif 'cert' in header:
                        consultant['certifications'] = value
                    elif 'contact' in header:
                        consultant['contact'] = value
                    elif 'email' in header:
                        consultant['email'] = value
                    elif 'website' in header:
                        consultant['website'] = value
                    elif 'project' in header or 'projekt' in header:
                        consultant['projects'] = value
                    elif 'service' in header:
                        consultant['services'] = value
                    elif 'team' in header:
                        consultant['teamSize'] = value
                    else:
                        consultant[header.replace(' ', '_')] = value

            if consultant.get('name'):
                consultants.append(consultant)

    return consultants

def extract_case_studies(doc_path):
    """Extract hotel case studies from document 4"""
    doc = Document(doc_path)
    case_studies = []
    current_case = {}

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        # Detect new case study (hotel name, typically bold)
        if para.runs and para.runs[0].bold and len(text) < 100:
            if current_case.get('name'):
                case_studies.append(current_case)
            current_case = {'name': text}
        else:
            if ':' in text:
                parts = text.split(':', 1)
                key = parts[0].strip().lower()
                value = parts[1].strip() if len(parts) > 1 else ""

                field_mapping = {
                    'location': 'location',
                    'city': 'location',
                    'size': 'size',
                    'storlek': 'size',
                    'budget': 'budget',
                    'circular percentage': 'circularPercentage',
                    'cirkulär procent': 'circularPercentage',
                    'reuse percentage': 'circularPercentage',
                    'återanvändning': 'circularPercentage',
                    'materials': 'materials',
                    'material': 'materials',
                    'challenges': 'challenges',
                    'utmaningar': 'challenges',
                    'outcomes': 'outcomes',
                    'resultat': 'outcomes',
                    'results': 'outcomes',
                    'images': 'images',
                    'bilder': 'images',
                    'references': 'references',
                    'referenser': 'references',
                    'architect': 'architect',
                    'arkitekt': 'architect',
                    'contractor': 'contractor',
                    'entreprenör': 'contractor',
                    'year': 'year',
                    'år': 'year',
                    'timeline': 'timeline',
                    'tidslinje': 'timeline',
                    'rooms': 'rooms',
                    'rum': 'rooms',
                    'savings': 'savings',
                    'besparingar': 'savings',
                    'description': 'description',
                    'beskrivning': 'description',
                }

                for key_pattern, field_name in field_mapping.items():
                    if key_pattern in key:
                        current_case[field_name] = value
                        break

    if current_case.get('name'):
        case_studies.append(current_case)

    # Extract from tables
    for table in doc.tables:
        headers = [clean_text(cell.text).lower() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            case_study = {}
            for idx, cell in enumerate(row.cells):
                value = clean_text(cell.text)
                if idx < len(headers) and headers[idx]:
                    header = headers[idx]
                    if 'hotel' in header or 'name' in header or 'namn' in header:
                        case_study['name'] = value
                    elif 'location' in header or 'stad' in header:
                        case_study['location'] = value
                    elif 'size' in header or 'storlek' in header:
                        case_study['size'] = value
                    elif 'budget' in header:
                        case_study['budget'] = value
                    elif 'circular' in header or 'reuse' in header or 'återanvänd' in header:
                        case_study['circularPercentage'] = value
                    elif 'material' in header:
                        case_study['materials'] = value
                    elif 'challenge' in header or 'utmaning' in header:
                        case_study['challenges'] = value
                    elif 'outcome' in header or 'result' in header:
                        case_study['outcomes'] = value
                    elif 'year' in header or 'år' in header:
                        case_study['year'] = value
                    elif 'room' in header or 'rum' in header:
                        case_study['rooms'] = value
                    elif 'saving' in header or 'besparing' in header:
                        case_study['savings'] = value
                    else:
                        case_study[header.replace(' ', '_')] = value

            if case_study.get('name'):
                case_studies.append(case_study)

    return case_studies

def extract_regulations(doc_path):
    """Extract Swedish regulatory requirements from document 5"""
    doc = Document(doc_path)
    regulations = {
        'fireSafety': [],
        'buildingCodes': [],
        'bvb': [],
        'materialStandards': [],
        'documentation': [],
        'other': []
    }

    current_category = 'other'
    current_item = {}

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        # Detect category headers
        text_lower = text.lower()
        if 'fire safety' in text_lower or 'brandsäkerhet' in text_lower:
            current_category = 'fireSafety'
            if current_item:
                regulations[current_category].append(current_item)
                current_item = {}
        elif 'building code' in text_lower or 'byggbestämmelser' in text_lower or 'boverket' in text_lower:
            current_category = 'buildingCodes'
            if current_item:
                regulations[current_category].append(current_item)
                current_item = {}
        elif 'bvb' in text_lower:
            current_category = 'bvb'
            if current_item:
                regulations[current_category].append(current_item)
                current_item = {}
        elif 'material standard' in text_lower or 'materialstandard' in text_lower:
            current_category = 'materialStandards'
            if current_item:
                regulations[current_category].append(current_item)
                current_item = {}
        elif 'documentation' in text_lower or 'dokumentation' in text_lower:
            current_category = 'documentation'
            if current_item:
                regulations[current_category].append(current_item)
                current_item = {}

        # Parse regulation items
        if para.runs and para.runs[0].bold:
            if current_item:
                regulations[current_category].append(current_item)
            current_item = {'title': text, 'details': []}
        else:
            if current_item:
                current_item['details'].append(text)
            else:
                current_item = {'title': text, 'details': []}

    if current_item:
        regulations[current_category].append(current_item)

    # Extract from tables
    for table in doc.tables:
        headers = [clean_text(cell.text).lower() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            reg_item = {}
            for idx, cell in enumerate(row.cells):
                value = clean_text(cell.text)
                if idx < len(headers) and headers[idx]:
                    header = headers[idx]
                    reg_item[header.replace(' ', '_')] = value

            # Try to categorize
            item_text = str(reg_item).lower()
            if 'fire' in item_text or 'brand' in item_text:
                regulations['fireSafety'].append(reg_item)
            elif 'bvb' in item_text:
                regulations['bvb'].append(reg_item)
            elif 'building' in item_text or 'bygg' in item_text:
                regulations['buildingCodes'].append(reg_item)
            elif 'material' in item_text:
                regulations['materialStandards'].append(reg_item)
            elif 'document' in item_text or 'dokument' in item_text:
                regulations['documentation'].append(reg_item)
            else:
                regulations['other'].append(reg_item)

    return regulations

def extract_implementation(doc_path):
    """Extract strategic roadmap and implementation details from document 6"""
    doc = Document(doc_path)
    implementation = {
        'roadmap': [],
        'marketAnalysis': {},
        'recommendations': [],
        'timeline': [],
        'resources': [],
        'risks': [],
        'kpis': []
    }

    current_section = 'roadmap'
    current_item = {}

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        # Detect section headers
        text_lower = text.lower()
        if 'roadmap' in text_lower or 'färdplan' in text_lower:
            current_section = 'roadmap'
        elif 'market analysis' in text_lower or 'marknadsanalys' in text_lower:
            current_section = 'marketAnalysis'
        elif 'recommendation' in text_lower or 'rekommendation' in text_lower:
            current_section = 'recommendations'
        elif 'timeline' in text_lower or 'tidslinje' in text_lower:
            current_section = 'timeline'
        elif 'resource' in text_lower or 'resurs' in text_lower:
            current_section = 'resources'
        elif 'risk' in text_lower:
            current_section = 'risks'
        elif 'kpi' in text_lower or 'metrics' in text_lower or 'mätetal' in text_lower:
            current_section = 'kpis'

        # Store content
        if para.runs and para.runs[0].bold and len(text) < 200:
            if current_item:
                if isinstance(implementation[current_section], list):
                    implementation[current_section].append(current_item)
            current_item = {'title': text, 'content': []}
        else:
            if current_item:
                if isinstance(current_item, dict) and 'content' in current_item:
                    current_item['content'].append(text)
                else:
                    current_item = {'title': text, 'content': []}
            else:
                if isinstance(implementation[current_section], list):
                    implementation[current_section].append({'text': text})
                elif isinstance(implementation[current_section], dict):
                    implementation[current_section][len(implementation[current_section])] = text

    if current_item:
        if isinstance(implementation[current_section], list):
            implementation[current_section].append(current_item)

    # Extract from tables
    for table in doc.tables:
        headers = [clean_text(cell.text).lower() for cell in table.rows[0].cells]
        table_data = []
        for row in table.rows[1:]:
            row_data = {}
            for idx, cell in enumerate(row.cells):
                value = clean_text(cell.text)
                if idx < len(headers) and headers[idx]:
                    row_data[headers[idx].replace(' ', '_')] = value
            if row_data:
                table_data.append(row_data)

        # Try to categorize table data
        if table_data:
            header_text = ' '.join(headers).lower()
            if 'roadmap' in header_text or 'phase' in header_text:
                implementation['roadmap'].extend(table_data)
            elif 'timeline' in header_text or 'schedule' in header_text:
                implementation['timeline'].extend(table_data)
            elif 'resource' in header_text:
                implementation['resources'].extend(table_data)
            elif 'risk' in header_text:
                implementation['risks'].extend(table_data)
            elif 'kpi' in header_text or 'metric' in header_text:
                implementation['kpis'].extend(table_data)
            else:
                implementation['recommendations'].extend(table_data)

    return implementation

def main():
    """Main extraction process"""
    base_path = Path('/Users/gabrielboen/Downloads/drive-download-20251108T110451Z-1-001')

    print("Starting data extraction...")

    # Document 1 & 2: B2B Reuse Operators
    print("\n1. Extracting B2B Reuse Operators...")
    suppliers1 = extract_suppliers(base_path / '1. COMPREHENSIVE B2B REUSE OPERATORS EXTRACTION.docx')
    suppliers2 = extract_suppliers(base_path / '2. FYRA B2B OPERATOR CAPABILITY MATRIX.docx')

    # Merge and deduplicate suppliers
    all_suppliers = suppliers1 + suppliers2
    unique_suppliers = []
    seen_names = set()
    for supplier in all_suppliers:
        name = supplier.get('name', '').lower()
        if name and name not in seen_names:
            seen_names.add(name)
            unique_suppliers.append(supplier)
        elif name and name in seen_names:
            # Merge data for duplicate entries
            for existing in unique_suppliers:
                if existing.get('name', '').lower() == name:
                    for key, value in supplier.items():
                        if key not in existing or not existing[key]:
                            existing[key] = value

    with open(base_path / 'suppliers.json', 'w', encoding='utf-8') as f:
        json.dump(unique_suppliers, f, indent=2, ensure_ascii=False)
    print(f"   ✓ Extracted {len(unique_suppliers)} suppliers → suppliers.json")

    # Document 3: Consultants
    print("\n2. Extracting Project Managers & Technical Consultants...")
    consultants = extract_consultants(base_path / '3. SWEDISH PROJECT MANAGERS & TECHNICAL CONSULTANTS FOR WHOLE-PROJECT HOSPITALITY RENOVATION.docx')

    with open(base_path / 'consultants.json', 'w', encoding='utf-8') as f:
        json.dump(consultants, f, indent=2, ensure_ascii=False)
    print(f"   ✓ Extracted {len(consultants)} consultants → consultants.json")

    # Document 4: Case Studies
    print("\n3. Extracting Hotel Case Studies...")
    case_studies = extract_case_studies(base_path / '4. FRONTRUNNER HOTELS - COMPLETE EXTRACTION.docx')

    with open(base_path / 'caseStudies.json', 'w', encoding='utf-8') as f:
        json.dump(case_studies, f, indent=2, ensure_ascii=False)
    print(f"   ✓ Extracted {len(case_studies)} case studies → caseStudies.json")

    # Document 5: Regulations
    print("\n4. Extracting Swedish Regulatory Requirements...")
    regulations = extract_regulations(base_path / '5. PRACTICAL REGULATORY GUIDE_ REUSED MATERIALS IN SWEDISH HOTEL RENOVATIONS.docx')

    with open(base_path / 'regulations.json', 'w', encoding='utf-8') as f:
        json.dump(regulations, f, indent=2, ensure_ascii=False)

    total_regs = sum(len(v) if isinstance(v, list) else 0 for v in regulations.values())
    print(f"   ✓ Extracted {total_regs} regulations across {len(regulations)} categories → regulations.json")

    # Document 6: Implementation
    print("\n5. Extracting Strategic Implementation Roadmap...")
    implementation = extract_implementation(base_path / '6. STRATEGIC IMPLEMENTATION ROADMAP.docx')

    with open(base_path / 'implementation.json', 'w', encoding='utf-8') as f:
        json.dump(implementation, f, indent=2, ensure_ascii=False)
    print(f"   ✓ Extracted implementation roadmap → implementation.json")

    # Create summary
    summary = {
        'extractionDate': '2025-11-08',
        'sourceDocuments': 6,
        'outputFiles': 5,
        'dataExtracted': {
            'suppliers': len(unique_suppliers),
            'consultants': len(consultants),
            'caseStudies': len(case_studies),
            'regulations': total_regs,
            'implementationSections': len(implementation)
        },
        'files': {
            'suppliers': 'suppliers.json',
            'consultants': 'consultants.json',
            'caseStudies': 'caseStudies.json',
            'regulations': 'regulations.json',
            'implementation': 'implementation.json'
        }
    }

    with open(base_path / 'extraction_summary.json', 'w', encoding='utf-8') as f:
        json.dump(summary, f, indent=2, ensure_ascii=False)

    print("\n" + "="*60)
    print("EXTRACTION COMPLETE!")
    print("="*60)
    print(f"\nTotal data extracted:")
    print(f"  • {len(unique_suppliers)} B2B Reuse Operators")
    print(f"  • {len(consultants)} Project Managers/Consultants")
    print(f"  • {len(case_studies)} Hotel Case Studies")
    print(f"  • {total_regs} Regulatory Requirements")
    print(f"  • Implementation Roadmap & Market Analysis")
    print(f"\nJSON files created:")
    print(f"  1. suppliers.json")
    print(f"  2. consultants.json")
    print(f"  3. caseStudies.json")
    print(f"  4. regulations.json")
    print(f"  5. implementation.json")
    print(f"  6. extraction_summary.json")
    print("\nAll files ready for Next.js integration!")

if __name__ == '__main__':
    main()
